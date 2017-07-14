# -*- coding: utf-8 -*-
#-------------------------------------------------------------------------------
# Name:        colorgrade_comparisond.py
# Purpose:     code for colorgrade analytics research
#               comparison between interpolation of dense data with filtered data with comparison statistics
#               this is just for research purpose and not for the final library
# Author:      claudio piccinini
#
# Created:     04/04/2017
#-------------------------------------------------------------------------------

#import time

import itertools
import math
import random
import shutil
#import json
#import shutil
import warnings
import sys
import os
import os.path
import re
import io

import selenium.webdriver as webdriver
import PIL.Image as Image

import numpy as np
#from scipy.interpolate import Rbf
#from scipy.misc import imsave
import pandas as pd

from osgeo import gdal
#from osgeo import gdalconst as gdct
#from osgeo import ogr
from osgeo import osr

import scipy.spatial.ckdtree as kdtree
#from scipy.spatial.distance import pdist
from scipy.spatial import ConvexHull

import matplotlib
matplotlib.use('Agg')
#import matplotlib.pyplot as plt
#from matplotlib import cm

from sklearn import linear_model

from bokeh.plotting import figure
from bokeh.io import output_file, save
#from bokeh.models import Label
from bokeh.models import NumeralTickFormatter, Label
from bokeh.palettes import d3
from bokeh.io import export_png  #new in veersion 0.12.6


#from rasterstats import point_query

from docx import Document
from docx.shared import Cm

from pysdss.utility import fileIO
from pysdss.utility import gdalRasterIO as gdalIO

from pysdss import utils
from pysdss.database import database as dt
from pysdss.filtering import filter
from pysdss.gridding import gridding as grd
from pysdss.utility import utils as utils2

import  pysdss.geostat.variance as vrc
import pysdss.geostat.fit as fit
from  pysdss.utils import rmse
from pysdss.geostat import kriging as krig


###some harddoded variables for database connection and database query
connstring = {"dbname": "grapes", "user": "claudio", "password": "claudio", "port": "5432", "host": "127.0.0.1"}
datasetid = 26
boundid = 1
#######################################################



def outinfo(outpath,
                    name,
                    filter_type,
                    interpolation,
                    first_last,
                    steps,
                    colorgrades,
                    counts,
                    variogram=""
            ):

    """

    Output a json file with the interpolation info (which will be used later to label charts)

    {
    "name": "moving average" | "inverse distance" | "simple kriging" | "ordinary kriging"
    "filter_type": "ordered" | "random"
    "interpolation": "free" | "onlyrows"     (all neighbours or just vineyard neighbours)
    "first_last": "true" | "false"    (take first/last point of a row)
    "steps":[2,3,4,....]
    "colorgrades":['a','b','c','d']
    "counts":['visible','raw']
    "variogram": "" | "global" | "local" | "byrow"     ("byrow" is a made with the row points only)
    }

    :param outpath:
    :param kwargs: the keys for the dictionary
    :return:
    """

    info = {
        "name":name,
        "filter_type":filter_type,
        "interpolation":interpolation,
        "first_last":first_last,
        "steps":steps,
        "colorgrades":colorgrades,
        "counts":counts,
        "variogram":variogram
    }
    utils.json_io(outpath, info)

def outputinfotext(paths, outdir):
    """
    Ouyput 'fullinfos.csv' with infos about the interpolation folders
    the structure of the file will be like:
    id	            name	        filter_type	 interpolation	 first_last	variogram	steps	        counts	            colorgrades
    a00bca573171a4f	simple kriging	random	    free	        TRUE	    global	    [5, 6, 7, 8]	['visible']	        ['a', 'b', 'c', 'd']
    a01b2eada44b842	moving average	random	    free	        FALSE		            [2, 3, 4, 5]	['visible', 'raw']	['a', 'b', 'c', 'd']

    :param paths: list of folder with interpolation results
    :param outdir: the output directory for the text file
    :return:
    """

    out = open(outdir + "/fullinfos.csv", "w")

    out.write("id, name, filter_type, interpolation, first_last,variogram,steps,counts,colorgrades\n")

    for p in paths:

        info = utils.json_io(p + "/info.json", direction="in")
        b = os.path.basename(p)
        id = b if b.strip() != '' else p.split('/')[-2]
        out.write(
        id+"," + info["name"]+ "," + info["filter_type"]+"," + info["interpolation"]+"," + \
        str(info["first_last"]) + "," + info["variogram"]+",\"" + repr(info["steps"])+"\",\"" + \
        repr(info["counts"])+ "\",\"" + repr(info["colorgrades"]) +"\"\n")


########## download generic 2 joined data table
def dowload_all_data_test(id, path):
    """
    download data for all 4 grades
    :param id:
    :param path:
    :return:
    """
    leftclm = ["row", "id_sdata", "lat", "lon"]
    rightclmn = ["a", "b", "c", "d", "keepa::int", "keepb::int", "keepc::int", "keepd::int"]

    outfile = path + id + ".csv"
    # complete dataset
    dt.download_data("id_sdata", "data.SensoDatum", "id_sensor", leftclm, "id_sdata", "data.colorgrade", rightclmn,
                     datasetid, outfile, orderclms=["row", "id_sdata"], conndict=connstring)
    name = utils.create_uniqueid()
    outfile = "/vagrant/code/pysdss/data/output/" + name + ".csv"
    # clipped dataset, result ordered
    dt.download_data("id_sdata", "data.SensoDatum", "id_sensor", leftclm, "id_sdata", "data.colorgrade", rightclmn,
                     datasetid, outfile, leftboundgeom="geom", boundtable="data.boundary", righttboundgeom="geom",
                     boundid="id_bound",
                     boundvalue=1, orderclms=["row", "id_sdata"], conndict=connstring)

def test_interpolate_and_filter():
    """

    :return:
    """

    # 1 download filtered data with the chosen and create a dataframe

    folder = "/vagrant/code/pysdss/data/output/text/"
    experimentfolder = "/vagrant/code/pysdss/experiments/interpolation/"

    id = utils.create_uniqueid()

    # create a directory to store settings
    # if not os.path.exists(experimentfolder + str(id)):
    #   os.mkdir(experimentfolder + str(id))


    # dowload_data_test(id, folder)
    dowload_all_data_test(id, folder)

    df = pd.read_csv(folder + id + ".csv")
    # strip and lower column names
    df.columns = df.columns.str.strip()
    df.columns = df.columns.str.lower()

    # here should be necessary some mapping to get the column names


    # 2 filter data and save to disk . this is the "truth" so we just delete the zeros
    # 1/03/17  we don't filter
    # keep, discard = filter.filter_byvalue(df, 0, ">", colname="d")
    keep = df

    # 3 save filtered data to disk after reprojection
    # reproject latitude, longitude #WGS 84 / UTM zone 11N
    west, north = utils2.reproject_coordinates(keep, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(keep, west, "lon")
    filter.set_field_value(keep, north, "lat")
    filter.set_field_value(keep, 0, "keepa")
    filter.set_field_value(keep, 0, "keepb")
    filter.set_field_value(keep, 0, "keepc")
    filter.set_field_value(keep, 0, "keepd")
    keep.to_csv(folder + id + "_keep.csv", index=False)

    '''west,north =utils.reproject_coordinates(discard, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(discard, west,"lon")
    filter.set_field_value(discard, north,"lat")
    filter.set_field_value(discard, 0, "keepd")
    discard.to_csv(folder + id +"_discard.csv",index=False)'''

    # 4 create the virtualdataset header for the datafile

    xml = '''<OGRVRTDataSource>
                <OGRVRTLayer name="{layername}">
                    <SrcDataSource>{fullname}</SrcDataSource>
                    <GeometryType>wkbPoint</GeometryType>
                    <GeometryField encoding="PointFromColumns" x="{easting}" y="{northing}" z="{elevation}"/>
                </OGRVRTLayer>
    </OGRVRTDataSource>
    '''

    # output the 4 virtual files for the 4 columns
    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
                "elevation": clr}
        newxml = xml.format(**data)
        f = open(folder + id + "_keep" + clr + ".vrt", "w")
        f.write(newxml)
        f.close()

    # 5 define an interpolation grid (the user can choose an existing grid or draw on screen)
    # here i am doing manually for testing, i just take the bounding box

    minx = math.floor(keep['lon'].min())
    maxx = math.ceil(keep['lon'].max())
    miny = math.floor(keep['lat'].min())
    maxy = math.ceil(keep['lat'].max())
    delty = maxy - miny
    deltx = maxx - minx

    # 6 log a running process to the database
    # 7 save grid metadata and interpolation metadata to the database historic tables
    # 8 execute gridding with gdal_array (set nodata value to zero or other value)


    algo = {
        "name": "average",
        "radius1": "0.4",
        "radius2": "0.4",
        "angle": "0",
        "min_points": "1",
        "nodata": "0"
    }
    params = {
        "-ot": "Float32",
        "-of": "GTiff",
        "-co": "",
        "-zfield": "",
        "-z_increase": "",
        "-z_multiply": "",
        "-a_srs": "EPSG:32611",
        "-spat": "",
        "-clipsrc": "",
        "-clipsrcsql": "",
        "-clipsrclayer": "",
        "-clipsrcwhere": "",
        "-l": "",
        "-where": "",
        "-sql": "",
        "-txe": str(minx) + " " + str(maxx),
        "-tye": str(miny) + " " + str(maxy),
        # pixel is 0.5 meters
        "-outsize": str(deltx * 2) + " " + str(delty * 2),
        "-a": algo,
        "-q": "",
        "src_datasource": "",
        "dst_filename": ""
    }

    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:

        params["src_datasource"] = folder + id + "_keep" + clr + ".vrt"
        params["dst_filename"] = folder + id + "_" + clr + ".tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        # call gdal_grid
        print("Getting the 'truth' raster for color " + clr)
        out, err = utils.run_tool(text)
        print("output" + out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    # 9 add result url to the database table
    # 10 log completed process (maybe should add the location of the result


    # 11 define the indexed array for the image

    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        d = gdal.Open(folder + id + "_" + clr + ".tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                              nodata=0)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id + "_indexedarray", (index, r_indexed, properties))
        d = None

    ######## rasterize rows with nearest neighbour

    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "row"}
    newxml = xml.format(**data)
    f = open(folder + id + "_keep_row.vrt", "w")
    f.write(newxml)
    f.close()

    algo = {
        "name": "nearest",
        "radius1": "0.4",
        "radius2": "0.4",
        "nodata": "0"
    }
    params = {
        "-ot": "Int16",
        "-of": "GTiff",
        "-co": "",
        "-zfield": "",
        "-z_increase": "",
        "-z_multiply": "",
        "-a_srs": "EPSG:32611",
        "-spat": "",
        "-clipsrc": "",
        "-clipsrcsql": "",
        "-clipsrclayer": "",
        "-clipsrcwhere": "",
        "-l": "",
        "-where": "",
        "-sql": "",
        "-txe": str(minx) + " " + str(maxx),
        "-tye": str(miny) + " " + str(maxy),
        "-outsize": str(deltx * 2) + " " + str(delty * 2),
        "-a": algo,
        "-q": "",
        "src_datasource": folder + id + "_keep_row.vrt",
        "dst_filename": folder + id + "_row.tiff"
    }

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the row raster")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # upload to numpy
    d = gdal.Open(folder + id + "_row.tiff")
    band = d.GetRasterBand(1)
    # apply index
    row_indexed = gdalIO.apply_index_to_single_band(band, index)
    d = None

    # collect statistics for the comparisons
    stats = {}
    mean = []
    std = []

    # define new rasters increasing the data filtering and interpolating TODO: create configuration files for automation
    for step in [2, 3, 4]:

        # add dictionary for this step
        stats[step] = {}

        # filter data
        k, d = filter.filter_bystep(keep, step=step)

        '''
        i already projected before
        west, north = utils.reproject_coordinates(k, "epsg:32611", "lon", "lat", False)
        filter.set_field_value(k, west, "lon")
        filter.set_field_value(k, north, "lat")
        filter.set_field_value(k, 0, "keepa")
        filter.set_field_value(k, 0, "keepb")
        filter.set_field_value(k, 0, "keepc")
        filter.set_field_value(k, 0, "keepd")'''
        k.to_csv(folder + id + "_keep_step" + str(step) + ".csv", index=False)

        '''west, north = utils.reproject_coordinates(d, "epsg:32611", "lon", "lat", False)
        filter.set_field_value(d, west, "lon")
        filter.set_field_value(d, north, "lat")
        filter.set_field_value(d, 0, "keepa")
        filter.set_field_value(d, 0, "keepb")
        filter.set_field_value(d, 0, "keepc")
        filter.set_field_value(d, 0, "keepd")'''
        d.to_csv(folder + id + "_discard_step" + str(step) + ".csv", index=False)

        # for clr in ["a"]:
        for clr in ["a", "b", "c", "d"]:
            data = {"layername": id + "_keep_step" + str(step),
                    "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                    "elevation": clr}
            newxml = xml.format(**data)
            f = open(folder + id + "_keep" + clr + "_step" + str(step) + ".vrt", "w")
            f.write(newxml)
            f.close()

        # interpolate

        algo = {
            "name": "invdist",
            "radius1": str(0.5 * step),
            "radius2": str(0.5 * step),
            "angle": "0",
            "min_points": "1",
            "nodata": "0"
        }
        params = {
            "-ot": "Float32",
            "-of": "GTiff",
            "-co": "",
            "-zfield": "",
            "-z_increase": "",
            "-z_multiply": "",
            "-a_srs": "EPSG:32611",
            "-spat": "",
            "-clipsrc": "",
            "-clipsrcsql": "",
            "-clipsrclayer": "",
            "-clipsrcwhere": "",
            "-l": "",
            "-where": "",
            "-sql": "",
            "-txe": str(minx) + " " + str(maxx),
            "-tye": str(miny) + " " + str(maxy),
            "-outsize": str(deltx * 2) + " " + str(delty * 2),
            "-a": algo,
            "-q": "",
            "src_datasource": "",
            "dst_filename": ""
        }

        # for clr in ["a"]:
        for clr in ["a", "b", "c", "d"]:

            params["src_datasource"] = folder + id + "_keep" + clr + "_step" + str(step) + ".vrt"
            params["dst_filename"] = folder + id + "_" + clr + "_step" + str(step) + ".tiff"

            # build gdal_grid request
            text = grd.build_gdal_grid_string(params, outtext=False)
            print(text)
            text = ["gdal_grid"] + text
            print(text)

            # call gdal_grid
            print("Getting filtered raster by step " + str(step) + "for color grade " + clr)
            out, err = utils.run_tool(text)
            print("output" + out)
            if err:
                print("error:" + err)
                raise Exception("gdal grid failed")

            # upload to numpy
            d = gdal.Open(folder + id + "_" + clr + "_step" + str(step) + ".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed = gdalIO.apply_index_to_single_band(band, index)
            d = None

            # calculate average, std, total area (here a pixel is 0.25 m2,
            # totla area for pixel > 0, total area for pixels with 0 value


            stats[step][clr] = {}

            for i in np.unique(row_indexed):  # get the row numbers

                area = 0.25

                # add dictionary for this row
                stats[step][clr][i] = {}

                # get a mask for the current row
                mask = row_indexed == i
                # statistics for current row

                # r_indexed is 2d , while new_r_indexed and mask are 1d
                stats[step][clr][i]['truth'] = [r_indexed[0, :][mask].mean(), r_indexed[0, :][mask].std(),
                                                r_indexed[0, :][mask].shape[0] * area,
                                                np.count_nonzero(r_indexed[0, :][mask]) * area,
                                                r_indexed[0, :][mask][r_indexed[0, :][mask] == 0].shape[0] * area]
                stats[step][clr][i]['interpolated'] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                                                       new_r_indexed[mask].shape[0] * area,
                                                       np.count_nonzero(new_r_indexed[mask]) * area,
                                                       new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area]
                # compare = r_indexed[0,:][mask] - new_r_indexed[mask]
                # stats[clr][step][i]['comparison'] = [compare.mean(),compare.std()]
                pass
                # calculate average and variance for this surface

                # mean.append(compare.mean())
                # std.append(compare.std())

    # save as a boxplot

    # print(mean)
    # print(std)
    return id, stats
    # {2: {"a": {1: {'truth': {}, 'interpolated': {}}}}}
# id, stats = interpolate_and_filter()
# folder = "/vagrant/code/pysdss/data/output/text/"
# fileIO.save_object( folder + id + "_statistics", stats)

##########################################################
##########################################################

#######################################################
##############old berry workflow 1#####################
#######################################################
##############berry workflow 1  # the following functions are for comparisons between dense data and filtered data
#### data is hardcoded because these are not production functions

def berrycolor_workflow_1_original(steps=[2,3,4], interp="invdist"):
    """
    similar to interpolate_and_filter but here as input i am using the original csv file and I am considering the
    number of berries as well
    
    steps: the filtering steps to use
    interp: the interpolation method to use on the filtered data   "invdist" or "average"
    
    
    :return:
    """

    if interp not in ["invdist","average"]:
        raise ValueError("Interpolation should be 'invdist' or 'average'")

    jsonpath = os.path.join(os.path.dirname(__file__), '../../../..', "pysdss/experiments/interpolation/")
    jsonpath = os.path.normpath(jsonpath)


    # 1 download filtered data with the chosen and create a dataframe

    folder = "/vagrant/code/pysdss/data/output/text/"
    experimentfolder = "/vagrant/code/pysdss/experiments/interpolation/"

    # create a folder to store the output
    id = utils.create_uniqueid()
    if not os.path.exists(folder+id):
        os.mkdir(folder+id)
    folder = folder+id + "/"


    # create a directory to store settings
    #if not os.path.exists(experimentfolder + str(id)):
    #   os.mkdir(experimentfolder + str(id))
    #dowload_data_test(id, folder)



    # dowload the data #TODO: create function to read from the database
    df = pd.read_csv("/vagrant/code/pysdss/data/input/2016_06_17.csv", usecols=["%Dataset_id"," Row"," Raw_fruit_count",
        " Visible_fruit_per_m"," Latitude"," Longitude","Harvestable_A","Harvestable_B","Harvestable_C","Harvestable_D"])
    df.columns = ["id","row","raw_fruit_count","visible_fruit_per_m", "lat", "lon","a","b","c","d"]

    #here should be necessary some mapping to get the column names


    # 2 filter data and save to disk . this is the "truth" so we just delete the zeros
    # 1/03/17  we don't filter
    #keep, discard = filter.filter_byvalue(df, 0, ">", colname="d")
    keep = df


    # 3 save filtered data to disk after reprojection
    # reproject latitude, longitude #WGS 84 / UTM zone 11N
    west,north = utils2.reproject_coordinates(keep, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(keep, west, "lon")
    filter.set_field_value(keep, north, "lat")
    filter.set_field_value(keep, 0, "keepa")
    filter.set_field_value(keep, 0, "keepb")
    filter.set_field_value(keep, 0, "keepc")
    filter.set_field_value(keep, 0, "keepd")
    keep.to_csv(folder + id + "_keep.csv", index=False)


    '''west,north =utils.reproject_coordinates(discard, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(discard, west,"lon")
    filter.set_field_value(discard, north,"lat")
    filter.set_field_value(discard, 0, "keepd")
    discard.to_csv(folder + id +"_discard.csv",index=False)'''


    # 4 create the virtualdataset header for the datafile

    with open ( jsonpath + "/vrtmodel.txt") as f:
        xml = f.read()

    # output the 4 virtual files for the 4 columns
    #for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        data = {"layername": id +"_keep", "fullname": folder + id +"_keep.csv", "easting":"lon", "northing":"lat", "elevation":clr}
        utils2.make_vrt(xml, data,folder + id+"_keep"+clr+".vrt" )

    # output the virtual files for raw fruit count and visible fruit count
    data = {"layername": id +"_keep", "fullname": folder + id +"_keep.csv", "easting":"lon", "northing":"lat", "elevation":"raw_fruit_count"}
    utils2.make_vrt(xml, data, folder + id+"_keep_rawfruit.vrt")

    data = {"layername": id +"_keep", "fullname": folder + id +"_keep.csv", "easting":"lon", "northing":"lat", "elevation":"visible_fruit_per_m"}
    utils2.make_vrt(xml, data, folder + id+"_keep_visiblefruit.vrt")


    # 5 define an interpolation grid (the user can choose an existing grid or draw on screen)
    #here i am doing manually for testing, i just take the bounding box

    minx = math.floor(keep['lon'].min())
    maxx = math.ceil(keep['lon'].max())
    miny = math.floor(keep['lat'].min())
    maxy = math.ceil( keep['lat'].max())
    delty = maxy - miny #height
    deltx = maxx - minx #width

    area_multiplier = 2 #2 to double the resulution and halve the pixel size to 0.5 meters

    # 6 log a running process to the database
    # 7 save grid metadata and interpolation metadata to the database historic tables
    # 8 execute gridding with gdal_array (set nodata value to zero or other value)


    path = jsonpath + "/average.json"
    params = utils.json_io(path, direction="in")

    params["-txe"]= str(minx) + " " + str(maxx)
    params["-tye"]= str(miny) + " " + str(maxy)
    # pixel is 0.5 meters
    params["-outsize"]= str(deltx * area_multiplier) + " " + str(delty * area_multiplier)

    #for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:

        params["src_datasource"] = folder + id+"_keep"+clr+".vrt"
        params["dst_filename"] = folder + id + "_" +clr + ".tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        #call gdal_grid
        print("Getting the 'truth' raster for color " + clr)
        out,err = utils.run_tool(text)
        print("output"+out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    ##rasterize raw fruit and visible fruit

    params["src_datasource"] = folder + id + "_keep_rawfruit.vrt"
    params["dst_filename"] = folder + id + "_rawfruit.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the 'truth' raster for raw fruit count")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    params["src_datasource"] = folder + id + "_keep_visiblefruit.vrt"
    params["dst_filename"] = folder + id + "_visiblefruit.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the 'truth' raster for visible fruit per m")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # 9 add result url to the database table
    # 10 log completed process (maybe should add the location of the result


    # 11 define the indexed array for the images, I used the same interpolation method and parameters therefore the index
    #  is the same

    #for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        d = gdal.Open(folder + id+ "_" + clr + ".tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1,bandLocation="byrow", nodata=-1)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id+ "_indexedarray_"+clr, (index, r_indexed, properties))
        d = None

    # define the indexed array for the raw and visible fruit
    d = gdal.Open(folder + id + "_rawfruit.tiff")
    index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                          nodata=-1)
    print("saving indexed array to disk")
    fileIO.save_object(folder + id + "_indexedarray_rawfruit", (index, r_indexed, properties))
    d = None

    d = gdal.Open(folder + id + "_visiblefruit.tiff")
    index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                          nodata=-1)
    print("saving indexed array to disk")
    fileIO.save_object(folder + id + "_indexedarray_visiblefruit", (index, r_indexed, properties))
    d = None

    ######## rasterize rows with nearest neighbour

    data = {"layername": id +"_keep", "fullname": folder + id +"_keep.csv", "easting":"lon", "northing":"lat", "elevation":"row"}
    utils2.make_vrt(xml, data, folder + id+"_keep_row.vrt")
    '''newxml = xml.format(**data)
    f = open(folder + id+"_keep_row.vrt", "w")
    f.write(newxml)
    f.close()'''

    path = jsonpath + "/nearest.json"
    params = utils.json_io(path, direction="in")

    params["-txe"]= str(minx) + " " + str(maxx)
    params["-tye"]= str(miny) + " " + str(maxy)
    params["-outsize"]= str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
    params["src_datasource"]= folder + id + "_keep_row.vrt"
    params["dst_filename"]= folder + id + "_row.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    #call gdal_grid
    print("Getting the row raster")
    out,err = utils.run_tool(text)
    print("output"+out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # upload to numpy
    d = gdal.Open(folder + id+"_row.tiff")
    band = d.GetRasterBand(1)
    # apply index
    row_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count, all the indexes are the same
    d = None

    #collect statistics for the comparisons
    stats={}
    #mean=[]
    #std=[]

    ####### get the indexed array for raw count
    index_raw, r_indexed_raw, properties_raw = fileIO.load_object(folder + id + "_indexedarray_rawfruit")
    ####### get the indexed array for visible count
    index_visible, r_indexed_visible, properties_visible = fileIO.load_object(
        folder + id + "_indexedarray_visiblefruit")

    # define new rasters increasing the data filtering and interpolating TODO: create configuration files for automation
    for step in steps:

        # add dictionary for this step
        stats[step] = {}

        # filter data by step
        k, d = filter.filter_bystep(keep, step=step)

        '''
        i already projected before
        west, north = utils.reproject_coordinates(k, "epsg:32611", "lon", "lat", False)
        filter.set_field_value(k, west, "lon")
        filter.set_field_value(k, north, "lat")
        filter.set_field_value(k, 0, "keepa")
        filter.set_field_value(k, 0, "keepb")
        filter.set_field_value(k, 0, "keepc")
        filter.set_field_value(k, 0, "keepd")'''
        k.to_csv(folder + id + "_keep_step"+str(step)+".csv", index=False)

        '''west, north = utils.reproject_coordinates(d, "epsg:32611", "lon", "lat", False)
        filter.set_field_value(d, west, "lon")
        filter.set_field_value(d, north, "lat")
        filter.set_field_value(d, 0, "keepa")
        filter.set_field_value(d, 0, "keepb")
        filter.set_field_value(d, 0, "keepc")
        filter.set_field_value(d, 0, "keepd")'''
        d.to_csv(folder + id + "_discard_step"+str(step)+".csv", index=False)

        #for clr in ["a"]:
        for clr in ["a", "b", "c", "d"]:

            data = {"layername": id + "_keep_step"+str(step), "fullname": folder + id + "_keep_step"+str(step)+".csv", "easting": "lon", "northing": "lat",
                    "elevation": clr}
            utils2.make_vrt(xml, data, folder + id + "_keep"+clr+"_step"+str(step)+".vrt")
            '''newxml = xml.format(**data)
            f = open(folder + id + "_keep"+clr+"_step"+str(step)+".vrt", "w")
            f.write(newxml)
            f.close()'''

        data = {"layername": id + "_keep_step" + str(step),
                "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                "elevation": "raw_fruit_count"}
        utils2.make_vrt(xml, data, folder + id + "_keep_raw_step" + str(step) + ".vrt")
        '''newxml = xml.format(**data)
        f = open(folder + id + "_keep_raw_step" + str(step) + ".vrt", "w")
        f.write(newxml)
        f.close()'''

        data = {"layername": id + "_keep_step" + str(step),
                "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                "elevation": "visible_fruit_per_m"}
        utils2.make_vrt(xml, data, folder + id + "_keep_visible_step" + str(step) + ".vrt")
        '''newxml = xml.format(**data)
        f = open(folder + id + "_keep_visible_step" + str(step) + ".vrt", "w")
        f.write(newxml)
        f.close()'''

        # interpolate
        if interp== "invdist":
            path = jsonpath + "/invdist.json"
        if interp== "average":
            path = jsonpath + "/average.json"

        params = utils.json_io(path, direction="in")

        params["-txe"]= str(minx) + " " + str(maxx)
        params["-tye"]= str(miny) + " " + str(maxy)
        params["-outsize"]= str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
        params["-a"]["radius1"] = str(0.5 * step)
        params["-a"]["radius2"] = str(0.5 * step)


        for clr in ["raw", "visible"]:

            params["src_datasource"] = folder + id + "_keep_" + clr + "_step" + str(step) + ".vrt"
            params["dst_filename"] = folder + id + "_" + clr + "_step" + str(step) + ".tiff"

            print(params)
            # build gdal_grid request
            text = grd.build_gdal_grid_string(params, outtext=False)
            print(text)
            text = ["gdal_grid"] + text
            print(text)

            # call gdal_grid
            print("Getting filtered raster by step " + str(step) + "for color " + clr)
            out, err = utils.run_tool(text)
            print("output" + out)
            if err:
                print("error:" + err)
                raise Exception("gdal grid failed")


        # upload to numpy
        d = gdal.Open(folder + id + "_raw_step" + str(step) + ".tiff")
        band = d.GetRasterBand(1)
        # apply index
        new_r_indexed_raw = gdalIO.apply_index_to_single_band(band, index)# this is the last index from the fruit count
        d = None

        d = gdal.Open(folder + id + "_visible_step" + str(step) + ".tiff")
        band = d.GetRasterBand(1)
        # apply index
        new_r_indexed_visible = gdalIO.apply_index_to_single_band(band, index)# this is the last index from the fruit count
        d = None

        #for clr in ["a"]:
        for clr in ["a", "b", "c", "d"]:

            params["src_datasource"] = folder + id + "_keep" + clr + "_step" + str(step)+".vrt"
            params["dst_filename"] = folder + id + "_"+clr + "_step"+str(step)+".tiff"

            # build gdal_grid request
            text = grd.build_gdal_grid_string(params, outtext=False)
            print(text)
            text = ["gdal_grid"] + text
            print(text)

            # call gdal_grid
            print("Getting filtered raster by step " + str(step) + "for color grade " + clr)
            out, err = utils.run_tool(text)
            print("output" + out)
            if err:
                print("error:" + err)
                raise Exception("gdal grid failed")

            # upload to numpy
            d = gdal.Open(folder + id + "_"+clr + "_step"+str(step)+".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed = gdalIO.apply_index_to_single_band(band, index)# this is the last index from the fruit count
            d = None

            # calculate average, std, total area (here a pixel is 0.25 m2,
            # totla area for pixel > 0, total area for pixels with 0 value

            ####### get the indexed array for this colour grade

            #d = gdal.Open(folder + id + "_" + clr + ".tiff")
            #band = d.GetRasterBand(1)
            #r_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count
            #d = None
            index, r_indexed, properties= fileIO.load_object(folder + id + "_indexedarray_" + clr)


            stats[step][clr] = {}

            for i in np.unique(row_indexed): #get the row numbers

                area = math.pow(1/area_multiplier,2)

                #add dictionary for this row
                stats[step][clr][i] = {}

                #get a mask for the current row
                mask = row_indexed == i
                #statistics for current row

                #average, std, totalarea,total nonzero area, total zeroarea, total raw fruit count,
                # average raw fruit count, std raw fruit count ,total visible fruitxm, average visible fruitXm, std visible fruitXm

                #r_indexed is 2d , while new_r_indexed and mask are 1d
                stats[step][clr][i]['truth'] = [r_indexed[0,:][mask].mean(), r_indexed[0,:][mask].std(),
                                                r_indexed[0,:][mask].shape[0]*area, np.count_nonzero(r_indexed[0,:][mask])*area,
                                                r_indexed[0,:][mask][r_indexed[0,:][mask] == 0].shape[0]*area,
                                                r_indexed_raw[0,:][mask].sum(),
                                                r_indexed_raw[0,:][mask].mean(),
                                                r_indexed_raw[0,:][mask].std(),
                                                r_indexed_visible[0,:][mask].sum(),
                                                r_indexed_visible[0,:][mask].mean(),
                                                r_indexed_visible[0,:][mask].std()]

                '''
                stats[step][clr][i]['truth'] = [r_indexed[mask].mean(), r_indexed[mask].std(),
                                                r_indexed[mask].shape[0]*area, np.count_nonzero(r_indexed[mask])*area,
                                                r_indexed[mask][r_indexed[mask] == 0].shape[0]*area,
                                                r_indexed_raw[0,:][mask].sum(), r_indexed_raw[0,:][mask].mean(),r_indexed_raw[0,:][mask].std(),
                                                r_indexed_visible[0,:][mask].sum(),r_indexed_visible[0,:][mask].mean()], r_indexed_visible[0,:][mask].std()]'''

                stats[step][clr][i]['interpolated'] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                                                       new_r_indexed[mask].shape[0]*area, np.count_nonzero(new_r_indexed[mask])*area,
                                                       new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0]*area,
                                                       new_r_indexed_raw[mask].sum(),
                                                       new_r_indexed_raw[mask].mean(),
                                                       new_r_indexed_raw[mask].std(),
                                                       new_r_indexed_visible[mask].sum(),
                                                       new_r_indexed_visible[mask].mean(),
                                                       new_r_indexed_visible[mask].std()]

                #compare = r_indexed[0,:][mask] - new_r_indexed[mask]
                #stats[clr][step][i]['comparison'] = [compare.mean(),compare.std()]
                #pass
        # calculate average and variance for this surface

        #mean.append(compare.mean())
        #std.append(compare.std())

    #save as a boxplot

    #print(mean)
    #print(std)
    return id,stats
    #{2: {"a": {1: {'truth': {}, 'interpolated': {}}}}}




############################################   30 05 2017 ######################################

##added parameter that specify we want random filtered data for row  ( take first/last and then random in between the 2)
##added parameter to force inteerpolation for points in the same row only (for higher filters the interpolation radius will
##overlap with points belonging to adjacent rows

#so there are 8 possibilities for each interpolation method (distance and moving average, for simple/ordinary kriging use berrycolor_workflow_kriging)
# ordered filter, first/last, free interpolation
# ordered filter, no first/last, free interpolation
# ordered filter, first/last, interpolation by row only
# ordered filter, no first/last, interpolation by row only
# random filter, first/last, free interpolation
# random filter, no first/last, free interpolation
# random filter, first/last, interpolation by row only
# random filter, no first/last, interpolation by row only


#for random filter the radius is calculated with the max distance between points on the same row


#for the inverse distance at present the power is set to the default value 2, the interpolation should be repeated for
#different power values and then root mean square error should be calculated to decide the optimal radius
#actually it should be also possible to experiment different radius setting, therefore creating a matrix radius/power


#           pow1    pow2    pow3    ...
#   radius1  x       x       x
#   radius2  x       x       x      ...
#   radius3  x       x       x      ...
#   radius4  x       x       x      ...


def berrycolor_workflow_1(steps=[2, 3, 4], interp="invdist", random_filter=False, first_last=False, force_interpolation_by_row=False, rowdirection="x"):
    """
    similar to interpolate_and_filter but here as input i am using the original csv file and I am considering the
    number of berries as well

    steps: the filtering steps to use
    interp: the interpolation method to use on the filtered data   "invdist" or "average", kriging will be added later
    random filter:   true to take random points (e,g step2 takes 50% random points for each row, false use ordered filter
    (e.g. step2 takes 1 point every 2 points)
    first_last: true to always have the first and last point of a vineyard row
    force_interpolation : true if i want to be sure the points belong to the current row only
    :return:
    """
    area_multiplier = 2  # 2 to double the resulution and halve the pixel size to 0.5 meters

    if interp not in ["invdist", "average"]:
        raise ValueError("Interpolation should be 'invdist' or 'average'")

    jsonpath = os.path.join(os.path.dirname(__file__), '../../../..', "pysdss/experiments/interpolation/")
    jsonpath = os.path.normpath(jsonpath)

    # 1 download filtered data with the chosen and create a dataframe

    folder = "/vagrant/code/pysdss/data/output/text/"
    experimentfolder = "/vagrant/code/pysdss/experiments/interpolation/"

    # create a folder to store the output
    id = utils.create_uniqueid()
    if not os.path.exists(folder + id):
        os.mkdir(folder + id)
    folder = folder + id + "/"

    outinfo(folder+"info.json",
            "moving average" if interp=="average" else "inverse distance",
            "random" if random_filter else "ordered",
            "onlyrows" if force_interpolation_by_row else "free",
            first_last,
            steps,
            colorgrades=['a', 'b', 'c', 'd'],
            counts=["visible","raw"]
            )

    # create a directory to store settings
    # if not os.path.exists(experimentfolder + str(id)):
    #   os.mkdir(experimentfolder + str(id))
    # dowload_data_test(id, folder)

    # dowload the data #TODO: create function to read from the database
    df = pd.read_csv("/vagrant/code/pysdss/data/input/2016_06_17.csv",
                     usecols=["%Dataset_id", " Row", " Raw_fruit_count",
                              " Visible_fruit_per_m", " Latitude", " Longitude", "Harvestable_A", "Harvestable_B",
                              "Harvestable_C", "Harvestable_D"])
    df.columns = ["id", "row", "raw_fruit_count", "visible_fruit_per_m", "lat", "lon", "a", "b", "c", "d"]

    # here should be necessary some mapping to get the column names


    # 2 filter data (and save to disk .
    # 1/03/17  we don't filter
    # keep, discard = filter.filter_byvalue(df, 0, ">", colname="d")
    keep = df

    # 3 save filtered data to disk after reprojection
    # reproject latitude, longitude #WGS 84 / UTM zone 11N
    west, north = utils2.reproject_coordinates(keep, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(keep, west, "lon")
    filter.set_field_value(keep, north, "lat")
    filter.set_field_value(keep, 0, "keepa")
    filter.set_field_value(keep, 0, "keepb")
    filter.set_field_value(keep, 0, "keepc")
    filter.set_field_value(keep, 0, "keepd")


    ##############search and remove duplicates
    #keep = utils.remove_duplicates(keep, ["lat", "lon"], operation="mean")
    ########################################

    keep.to_csv(folder + id + "_keep.csv", index=False)

    '''west,north =utils.reproject_coordinates(discard, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(discard, west,"lon")
    filter.set_field_value(discard, north,"lat")
    filter.set_field_value(discard, 0, "keepd")
    discard.to_csv(folder + id +"_discard.csv",index=False)'''

    # 4 create the virtualdataset header for the datafile

    with open(jsonpath + "/vrtmodel.txt") as f:
        xml = f.read()

    # output the 4 virtual files for the 4 columns
    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
                "elevation": clr}
        utils2.make_vrt(xml, data, folder + id + "_keep" + clr + ".vrt")

    # output the virtual files for raw fruit count and visible fruit count
    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "raw_fruit_count"}
    utils2.make_vrt(xml, data, folder + id + "_keep_rawfruit.vrt")

    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "visible_fruit_per_m"}
    utils2.make_vrt(xml, data, folder + id + "_keep_visiblefruit.vrt")

    # 5 define an interpolation grid (the user can choose an existing grid or draw on screen)
    # here i am doing manually for testing, i just take the bounding box

    minx = math.floor(keep['lon'].min())
    maxx = math.ceil(keep['lon'].max())
    miny = math.floor(keep['lat'].min())
    maxy = math.ceil(keep['lat'].max())
    delty = maxy - miny  # height
    deltx = maxx - minx  # width



    # 6 log a running process to the database
    # 7 save grid metadata and interpolation metadata to the database historic tables
    # 8 execute gridding with gdal_array (set nodata value to zero or other value)


    path = jsonpath + "/average.json"
    params = utils.json_io(path, direction="in")

    params["-txe"] = str(minx) + " " + str(maxx)
    params["-tye"] = str(miny) + " " + str(maxy)
    # pixel is 0.5 meters
    params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)

    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:

        params["src_datasource"] = folder + id + "_keep" + clr + ".vrt"
        params["dst_filename"] = folder + id + "_" + clr + ".tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        # call gdal_grid
        print("Getting the 'truth' raster for color " + clr)
        out, err = utils.run_tool(text)
        print("output" + out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    ##rasterize raw fruit and visible fruit

    params["src_datasource"] = folder + id + "_keep_rawfruit.vrt"
    params["dst_filename"] = folder + id + "_rawfruit.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the 'truth' raster for raw fruit count")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    params["src_datasource"] = folder + id + "_keep_visiblefruit.vrt"
    params["dst_filename"] = folder + id + "_visiblefruit.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the 'truth' raster for visible fruit per m")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # 9 add result url to the database table
    # 10 log completed process (maybe should add the location of the result


    # 11 define the indexed array for the images, I used the same interpolation method and parameters therefore the index
    #  is the same

    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        d = gdal.Open(folder + id + "_" + clr + ".tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                              nodata=-1)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id + "_indexedarray_" + clr, (index, r_indexed, properties))
        d = None

    # define the indexed array for the raw and visible fruit
    d = gdal.Open(folder + id + "_rawfruit.tiff")
    index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                          nodata=-1)
    print("saving indexed array to disk")
    fileIO.save_object(folder + id + "_indexedarray_rawfruit", (index, r_indexed, properties))
    d = None

    d = gdal.Open(folder + id + "_visiblefruit.tiff")
    index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                          nodata=-1)
    print("saving indexed array to disk")
    fileIO.save_object(folder + id + "_indexedarray_visiblefruit", (index, r_indexed, properties))
    d = None

    ######## rasterize rows with nearest neighbour

    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "row"}
    utils2.make_vrt(xml, data, folder + id + "_keep_row.vrt")
    '''newxml = xml.format(**data)
    f = open(folder + id+"_keep_row.vrt", "w")
    f.write(newxml)
    f.close()'''

    path = jsonpath + "/nearest.json"
    params = utils.json_io(path, direction="in")

    params["-txe"] = str(minx) + " " + str(maxx)
    params["-tye"] = str(miny) + " " + str(maxy)
    params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
    params["src_datasource"] = folder + id + "_keep_row.vrt"
    params["dst_filename"] = folder + id + "_row.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the row raster")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # upload to numpy
    d = gdal.Open(folder + id + "_row.tiff")
    band = d.GetRasterBand(1)
    # apply index
    row_indexed = gdalIO.apply_index_to_single_band(band,
                                                    index)  # this is the last index from the fruit count, all the indexes are the same
    d = None

    # collect statistics for the comparisons
    stats = {}
    # mean=[]
    # std=[]

    ####### get the indexed array for raw count
    index_raw, r_indexed_raw, properties_raw = fileIO.load_object(folder + id + "_indexedarray_rawfruit")
    ####### get the indexed array for visible count
    index_visible, r_indexed_visible, properties_visible = fileIO.load_object(
        folder + id + "_indexedarray_visiblefruit")

    if force_interpolation_by_row:

        ## make temporary directory to store the interpolated vineyard rows
        #################################################
        tempfolder = folder + "/temprasters/"
        os.mkdir(tempfolder)

        # define new rasters increasing the data filtering and interpolating TODO: create configuration files for automation
        for step in steps:

            # filter data by step
            k, d = filter.filter_bystep(keep, step=step, rand=random_filter, first_last=first_last,rowname="row")
            k.to_csv(folder + id + "_keep_step" + str(step) + ".csv", index=False)
            d.to_csv(folder + id + "_discard_step" + str(step) + ".csv", index=False)

            # find unique row numbers
            rows = k["row"].unique()

            # add dictionary for this step
            stats[step] = {}
            for clr in ["a", "b", "c", "d"]:
                stats[step][clr] = {}
                for row in rows:
                    stats[step][clr][row] = {}
                    stats[step][clr][row]['truth'] = []
                    stats[step][clr][row]['interpolated'] = []

            #load xml with columns
            with open(jsonpath + "/vrtmodel_row.txt") as f:
                xml_row = f.read()

            # for clr in ["a"]:
            for clr in ["a", "b", "c", "d"]:
                data = {"layername": id + "_keep_step" + str(step),
                        "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon",
                        "northing": "lat",
                        "elevation": clr, "rowname": "row"}
                utils2.make_vrt(xml_row, data, folder + id + "_keep" + clr + "_step" + str(step) + ".vrt")


            data = {"layername": id + "_keep_step" + str(step),
                    "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                    "elevation": "raw_fruit_count", "rowname": "row"}
            utils2.make_vrt(xml_row, data, folder + id + "_keep_raw_step" + str(step) + ".vrt")


            data = {"layername": id + "_keep_step" + str(step),
                    "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                    "elevation": "visible_fruit_per_m", "rowname": "row"}
            utils2.make_vrt(xml_row, data, folder + id + "_keep_visible_step" + str(step) + ".vrt")

            # prepare interpolation parameters
            if interp == "invdist":
                path = jsonpath + "/invdist.json"
            if interp == "average":
                path = jsonpath + "/average.json"

            params = utils.json_io(path, direction="in")

            params["-txe"] = str(minx) + " " + str(maxx)
            params["-tye"] = str(miny) + " " + str(maxy)
            params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)


            for clr in ["raw", "visible"]:

                #for each row if random calculate the max distance for the current row
                #execute gdalgrid for jus the row points
                #calculate statistics for the row and fill in the output data structure


                for row in rows:
                    # select one vineyard row
                    rowdf = k[k["row"] == int(row)]

                    # if random_filter we will set the radius to the max distance between points in a row
                    if random_filter: max_point_distance = utils.max_point_distance(rowdf, "lon", "lat", "row",
                                                                                    direction=rowdirection,
                                                                                    rem_duplicates=False)

                    if random_filter:  # if we have random points we set we set the radius to the max distance between points plus a buffer of 0.5
                        params["-a"]["radius1"] = str(max_point_distance + 0.5)
                        params["-a"]["radius2"] = str(max_point_distance + 0.5)
                    else:
                        params["-a"]["radius1"] = str(0.5 * step)
                        params["-a"]["radius2"] = str(0.5 * step)

                    #this will interpolate only the current row data
                    params["-where"] = "row="+ str(int(row))

                    params["src_datasource"] = folder + id + "_keep_" + clr + "_step" + str(step) + ".vrt"
                    params["dst_filename"] = tempfolder + id + "_" + clr + "_step" + str(step) + "_row"+str(row)+".tiff"

                    print(params)

                    # build gdal_grid request
                    text = grd.build_gdal_grid_string(params, outtext=False)
                    print(text)
                    text = ["gdal_grid"] + text
                    print(text)

                    # call gdal_grid
                    print("Getting filtered raster by step " + str(step) + "for field " + clr +" and row" + str(row))
                    out, err = utils.run_tool(text)
                    print("output" + out)
                    if err:
                        print("error:" + err)
                        raise Exception("gdal grid failed")


            # extract statistics for raw and visible
            for clr in ["a", "b", "c", "d"]:
                for row in rows:
                    # upload to numpy
                    d = gdal.Open(tempfolder + id + "_raw_step" + str(step) + "_row"+str(row) + ".tiff")
                    band = d.GetRasterBand(1)
                    # apply index
                    new_r_indexed_raw = gdalIO.apply_index_to_single_band(band,
                                                                          index)  # this is the last index from the fruit count
                    d = None

                    d = gdal.Open(tempfolder + id + "_visible_step" + str(step) + "_row"+str(row) + ".tiff")
                    band = d.GetRasterBand(1)
                    # apply index
                    new_r_indexed_visible = gdalIO.apply_index_to_single_band(band,
                                                                              index)  # this is the last index from the fruit count
                    d = None

                    #update structure with partial data
                    # get a mask for the current row
                    mask = row_indexed == row
                    # statistics for current row
                    # average, std, totalarea,total nonzero area, total zeroarea, total raw fruit count,
                    # average raw fruit count, std raw fruit count ,total visible fruitxm, average visible fruitXm, std visible fruitXm

                    stats[step][clr][row]['truth'] = [None, None, None, None, None,
                                          r_indexed_raw[0, :][mask].sum(),
                                          r_indexed_raw[0, :][mask].mean(),
                                          r_indexed_raw[0, :][mask].std(),
                                          r_indexed_visible[0, :][mask].sum(),
                                          r_indexed_visible[0, :][mask].mean(),
                                          r_indexed_visible[0, :][mask].std()]

                    stats[step][clr][row]['interpolated'] = [None, None, None, None, None,
                                                 new_r_indexed_raw[mask].sum(),
                                                 new_r_indexed_raw[mask].mean(),
                                                 new_r_indexed_raw[mask].std(),
                                                 new_r_indexed_visible[mask].sum(),
                                                 new_r_indexed_visible[mask].mean(),
                                                 new_r_indexed_visible[mask].std()]

            for clr in ["a", "b", "c", "d"]:

                #for each row if random calculate the max distance for the current row
                #execute gdalgrid for jus the row points
                #calculate statistics for the row and fill in the output data structure

                for row in rows:
                    # select one vineyard row
                    rowdf = k[k["row"] == int(row)]

                    # if random_filter we will set the radius to the max distance between points in a row
                    if random_filter: max_point_distance = utils.max_point_distance(rowdf, "lon", "lat", "row",
                                                                                    direction=rowdirection,
                                                                                    rem_duplicates=False)

                    if random_filter:  # if we have random points we set we set the radius to the max distance between points plus a buffer of 0.5
                        params["-a"]["radius1"] = str(max_point_distance + 0.5)
                        params["-a"]["radius2"] = str(max_point_distance + 0.5)
                    else:
                        params["-a"]["radius1"] = str(0.5 * step)
                        params["-a"]["radius2"] = str(0.5 * step)

                    #this will interpolate only the current row data
                    params["-where"] = "row="+ str(int(row))

                    params["src_datasource"] = folder + id + "_keep" + clr + "_step" + str(step) + ".vrt"
                    params["dst_filename"] = tempfolder + id + "_" + clr + "_step" + str(step) + "_row"+str(row) + ".tiff"

                    # build gdal_grid request
                    text = grd.build_gdal_grid_string(params, outtext=False)
                    print(text)
                    text = ["gdal_grid"] + text
                    print(text)

                    # call gdal_grid
                    print("Getting filtered raster by step " + str(step) + "for color grade " + clr)
                    out, err = utils.run_tool(text)
                    print("output" + out)
                    if err:
                        print("error:" + err)
                        raise Exception("gdal grid failed")

                    # upload to numpy
                    d = gdal.Open(tempfolder + id + "_" + clr + "_step" + str(step) + "_row"+str(row) + ".tiff")
                    band = d.GetRasterBand(1)
                    # apply index
                    new_r_indexed = gdalIO.apply_index_to_single_band(band,
                                                                      index)  # this is the last index from the fruit count
                    d = None

                    # calculate average, std, total area (here a pixel is 0.25 m2,
                    # totla area for pixel > 0, total area for pixels with 0 value

                    ####### get the indexed array for this colour grade

                    # d = gdal.Open(folder + id + "_" + clr + ".tiff")
                    # band = d.GetRasterBand(1)
                    # r_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count
                    # d = None
                    index, r_indexed, properties = fileIO.load_object(folder + id + "_indexedarray_" + clr)

                    # update structure with partial data
                    area = math.pow(1 / area_multiplier, 2)
                    # get a mask for the current row
                    mask = row_indexed == row
                    #fill in with data
                    partial_truth_data = stats[step][clr][row]['truth']
                    partial_truth_data[:5] = [r_indexed[0, :][mask].mean(), r_indexed[0, :][mask].std(),
                    r_indexed[0, :][mask].shape[0] * area,
                    np.count_nonzero(r_indexed[0, :][mask]) * area,
                    r_indexed[0, :][mask][r_indexed[0, :][mask] == 0].shape[0] * area]

                    partial_interpolated_data = stats[step][clr][row]['interpolated']
                    partial_interpolated_data[:5] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                    new_r_indexed[mask].shape[0] * area,
                    np.count_nonzero(new_r_indexed[mask]) * area,
                    new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area]



    else: # no force interpolation by row

        # define new rasters increasing the data filtering and interpolating TODO: create configuration files for automation
        for step in steps:
            # add dictionary for this step
            stats[step] = {}

            # filter data by step
            k, d = filter.filter_bystep(keep, step=step, rand=random_filter, first_last=first_last,rowname="row")
            # if random_filter we will set the radius to the max distance between points along the rows
            if random_filter: max_point_distance = utils.max_point_distance(k, "lon", "lat", "row", direction=rowdirection)

            '''
            i already projected before
            west, north = utils.reproject_coordinates(k, "epsg:32611", "lon", "lat", False)
            filter.set_field_value(k, west, "lon")
            filter.set_field_value(k, north, "lat")
            filter.set_field_value(k, 0, "keepa")
            filter.set_field_value(k, 0, "keepb")
            filter.set_field_value(k, 0, "keepc")
            filter.set_field_value(k, 0, "keepd")'''
            k.to_csv(folder + id + "_keep_step" + str(step) + ".csv", index=False)

            '''west, north = utils.reproject_coordinates(d, "epsg:32611", "lon", "lat", False)
            filter.set_field_value(d, west, "lon")
            filter.set_field_value(d, north, "lat")
            filter.set_field_value(d, 0, "keepa")
            filter.set_field_value(d, 0, "keepb")
            filter.set_field_value(d, 0, "keepc")
            filter.set_field_value(d, 0, "keepd")'''
            d.to_csv(folder + id + "_discard_step" + str(step) + ".csv", index=False)

            # for clr in ["a"]:
            for clr in ["a", "b", "c", "d"]:
                data = {"layername": id + "_keep_step" + str(step),
                        "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                        "elevation": clr}
                utils2.make_vrt(xml, data, folder + id + "_keep" + clr + "_step" + str(step) + ".vrt")
                '''newxml = xml.format(**data)
                f = open(folder + id + "_keep"+clr+"_step"+str(step)+".vrt", "w")
                f.write(newxml)
                f.close()'''

            data = {"layername": id + "_keep_step" + str(step),
                    "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                    "elevation": "raw_fruit_count"}
            utils2.make_vrt(xml, data, folder + id + "_keep_raw_step" + str(step) + ".vrt")
            '''newxml = xml.format(**data)
            f = open(folder + id + "_keep_raw_step" + str(step) + ".vrt", "w")
            f.write(newxml)
            f.close()'''

            data = {"layername": id + "_keep_step" + str(step),
                    "fullname": folder + id + "_keep_step" + str(step) + ".csv", "easting": "lon", "northing": "lat",
                    "elevation": "visible_fruit_per_m"}
            utils2.make_vrt(xml, data, folder + id + "_keep_visible_step" + str(step) + ".vrt")
            '''newxml = xml.format(**data)
            f = open(folder + id + "_keep_visible_step" + str(step) + ".vrt", "w")
            f.write(newxml)
            f.close()'''

            # interpolate
            if interp == "invdist":
                path = jsonpath + "/invdist.json"
            if interp == "average":
                path = jsonpath + "/average.json"

            params = utils.json_io(path, direction="in")

            params["-txe"] = str(minx) + " " + str(maxx)
            params["-tye"] = str(miny) + " " + str(maxy)
            params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
            if random_filter: #if we have random points we set we set the radius to the max distance between points plus a buffer of 0.5
                params["-a"]["radius1"] = str(max_point_distance + 0.5)
                params["-a"]["radius2"] = str(max_point_distance + 0.5)
            else:
                params["-a"]["radius1"] = str(0.5 * step)
                params["-a"]["radius2"] = str(0.5 * step)

            for clr in ["raw", "visible"]:

                params["src_datasource"] = folder + id + "_keep_" + clr + "_step" + str(step) + ".vrt"
                params["dst_filename"] = folder + id + "_" + clr + "_step" + str(step) + ".tiff"

                print(params)
                # build gdal_grid request
                text = grd.build_gdal_grid_string(params, outtext=False)
                print(text)
                text = ["gdal_grid"] + text
                print(text)

                # call gdal_grid
                print("Getting filtered raster by step " + str(step) + "for color " + clr)
                out, err = utils.run_tool(text)
                print("output" + out)
                if err:
                    print("error:" + err)
                    raise Exception("gdal grid failed")

            # upload to numpy
            d = gdal.Open(folder + id + "_raw_step" + str(step) + ".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed_raw = gdalIO.apply_index_to_single_band(band,
                                                                  index)  # this is the last index from the fruit count
            d = None

            d = gdal.Open(folder + id + "_visible_step" + str(step) + ".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed_visible = gdalIO.apply_index_to_single_band(band,
                                                                      index)  # this is the last index from the fruit count
            d = None

            # for clr in ["a"]:
            for clr in ["a", "b", "c", "d"]:

                params["src_datasource"] = folder + id + "_keep" + clr + "_step" + str(step) + ".vrt"
                params["dst_filename"] = folder + id + "_" + clr + "_step" + str(step) + ".tiff"

                # build gdal_grid request
                text = grd.build_gdal_grid_string(params, outtext=False)
                print(text)
                text = ["gdal_grid"] + text
                print(text)

                # call gdal_grid
                print("Getting filtered raster by step " + str(step) + "for color grade " + clr)
                out, err = utils.run_tool(text)
                print("output" + out)
                if err:
                    print("error:" + err)
                    raise Exception("gdal grid failed")

                # upload to numpy
                d = gdal.Open(folder + id + "_" + clr + "_step" + str(step) + ".tiff")
                band = d.GetRasterBand(1)
                # apply index
                new_r_indexed = gdalIO.apply_index_to_single_band(band,
                                                                  index)  # this is the last index from the fruit count
                d = None

                # calculate average, std, total area (here a pixel is 0.25 m2,
                # totla area for pixel > 0, total area for pixels with 0 value

                ####### get the indexed array for this colour grade

                # d = gdal.Open(folder + id + "_" + clr + ".tiff")
                # band = d.GetRasterBand(1)
                # r_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count
                # d = None
                index, r_indexed, properties = fileIO.load_object(folder + id + "_indexedarray_" + clr)

                stats[step][clr] = {}

                for i in np.unique(row_indexed):  # get the row numbers

                    area = math.pow(1 / area_multiplier, 2)

                    # add dictionary for this row
                    stats[step][clr][i] = {}

                    # get a mask for the current row
                    mask = row_indexed == i
                    # statistics for current row

                    # average, std, totalarea,total nonzero area, total zeroarea, total raw fruit count,
                    # average raw fruit count, std raw fruit count ,total visible fruitxm, average visible fruitXm, std visible fruitXm

                    # r_indexed is 2d , while new_r_indexed and mask are 1d
                    stats[step][clr][i]['truth'] = [r_indexed[0, :][mask].mean(), r_indexed[0, :][mask].std(),
                                                    r_indexed[0, :][mask].shape[0] * area,
                                                    np.count_nonzero(r_indexed[0, :][mask]) * area,
                                                    r_indexed[0, :][mask][r_indexed[0, :][mask] == 0].shape[0] * area,
                                                    r_indexed_raw[0, :][mask].sum(),
                                                    r_indexed_raw[0, :][mask].mean(),
                                                    r_indexed_raw[0, :][mask].std(),
                                                    r_indexed_visible[0, :][mask].sum(),
                                                    r_indexed_visible[0, :][mask].mean(),
                                                    r_indexed_visible[0, :][mask].std()]

                    '''
                    stats[step][clr][i]['truth'] = [r_indexed[mask].mean(), r_indexed[mask].std(),
                                                    r_indexed[mask].shape[0]*area, np.count_nonzero(r_indexed[mask])*area,
                                                    r_indexed[mask][r_indexed[mask] == 0].shape[0]*area,
                                                    r_indexed_raw[0,:][mask].sum(), r_indexed_raw[0,:][mask].mean(),r_indexed_raw[0,:][mask].std(),
                                                    r_indexed_visible[0,:][mask].sum(),r_indexed_visible[0,:][mask].mean()], r_indexed_visible[0,:][mask].std()]'''

                    stats[step][clr][i]['interpolated'] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                                                           new_r_indexed[mask].shape[0] * area,
                                                           np.count_nonzero(new_r_indexed[mask]) * area,
                                                           new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area,
                                                           new_r_indexed_raw[mask].sum(),
                                                           new_r_indexed_raw[mask].mean(),
                                                           new_r_indexed_raw[mask].std(),
                                                           new_r_indexed_visible[mask].sum(),
                                                           new_r_indexed_visible[mask].mean(),
                                                           new_r_indexed_visible[mask].std()]

                    # compare = r_indexed[0,:][mask] - new_r_indexed[mask]
                    # stats[clr][step][i]['comparison'] = [compare.mean(),compare.std()]
                    # pass
                    # calculate average and variance for this surface

                    # mean.append(compare.mean())
                    # std.append(compare.std())


    # print(mean)
    # print(std)
    return id, stats
    # {2: {"a": {1: {'truth': {}, 'interpolated': {}}}}}


#################################################### kriging #################################################



def get_variogram(vardata, hh, maxlagdistance, outfile, forcenugget=True):
    """
    Calculate semivariance and fit the semivariogram
    :param vardata: x,y,value
    :param hh: bin size
    :param maxlagdistance: maxdistance for semivariance
    :param outfile: full path name to output file (without extension)
    :return: the semivariogram object
    """

    print("calculte empirical semivariance")

    # i set the hh to 5 neters and the max distance to the smallest field edge
    # hh = 2
    # lags = range(0,min(delty ,deltx), hh)
    # lags = range(0, int((delty+deltx)/2), hh)

    # set lags from 0 to the interpolation radius
    # lags = range(0, int(radius), hh)

    #lags = np.arange(0, int(maxlagdistance * 2), hh)
    lags = np.arange(0, int(maxlagdistance*1.25), hh)
    gamma = vrc.semivar2(vardata, lags, hh)
    # covariance = vrc.covar(data, lags, hh)

    # chart empirical semivariance and covariance
    output_file(outfile + ".html")
    f = figure()
    f.line(gamma[0, :], gamma[1, :], line_color="green", line_width=2, legend="Semivariance")
    f.square(gamma[0, :], gamma[1, :], fill_color=None, line_color="green", legend="Semivariance")
    # f.line(cov[0, :], cov[1, :], line_color="red", line_width=2, legend="Covariance")
    # f.square(cov[0, :], cov[1, :], fill_color=None, line_color="red", legend="Covariance")
    f.legend.location = "top_left"
    save(f)

    print("fit semivariogram")
    # choose the model with lowest rmse (we use sphrical and exponential)
    semivariograms = []
    semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.spherical, forcenugget=forcenugget))
    #semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.linear, forcenugget=forcenugget))
    #semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.gaussian, forcenugget=forcenugget))
    semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.exponential,forcenugget=forcenugget))
    rsmes = [rmse(gamma[0], i(gamma[0])) for i in semivariograms]
    semivariogram = semivariograms[rsmes.index(min(rsmes))]

    # chart the fitted models
    output_file(outfile + "_fitted.html")
    f = figure()
    f.circle(gamma[0], gamma[1], fill_color=None, line_color="blue")
    f.line(gamma[0], semivariogram(gamma[0]), line_color="red", line_width=2)
    f.legend.location = "top_left"
    save(f)

    fileIO.save_object(outfile + ".var", semivariogram, usedill=True)

    return semivariogram


def savegrids(grid, gridpath, grid_error, grid_errorpath, xsize, ysize, geotr, nodata, epsg):
    """
    Output interpolated grids
    :param grid:
    :param gridpath:
    :param grid_error:
    :param grid_errorpath:
    :param xsize:
    :param ysize:
    :param geotr:
    :param nodata:
    :param epsg:
    :return:
    """

    outRaster = None
    try:
        # rastergrid = folder + id + "_visible_step" + str(step) + "_" + interp + variogram + ".tiff"
        rastergrid = gridpath + ".tiff"

        driver = gdal.GetDriverByName('GTiff')
        outRaster = driver.Create(rastergrid, xsize, ysize, 1, gdal.GDT_Float32)
        outRaster.SetGeoTransform(geotr)
        outband = outRaster.GetRasterBand(1)
        outband.SetNoDataValue(nodata)
        # outband.WriteArray( np.flip(grid,0))
        outband.WriteArray(grid)
        outRasterSRS = osr.SpatialReference()
        outRasterSRS.ImportFromEPSG(int(epsg))
        outRaster.SetProjection(outRasterSRS.ExportToWkt())
        outband.FlushCache()
    finally:
        if outRaster: outRaster = None
    try:
        # rastergriderror = folder + id + "_visible_step" + str(step) + "_" + interp + variogram + "_error.tiff"
        rastergriderror = grid_errorpath + ".tiff"

        driver = gdal.GetDriverByName('GTiff')
        outRaster = driver.Create(rastergriderror, xsize,
                                  ysize, 1, gdal.GDT_Float32)
        outRaster.SetGeoTransform(geotr)
        outband = outRaster.GetRasterBand(1)
        outband.SetNoDataValue(nodata)
        # outband.WriteArray( np.flip(grid_error,0))
        outband.WriteArray(grid_error)
        outRasterSRS = osr.SpatialReference()
        outRasterSRS.ImportFromEPSG(int(epsg))
        outRaster.SetProjection(outRasterSRS.ExportToWkt())
        outband.FlushCache()
    finally:
        if outRaster: outRaster = None


def fillin_grids_counts(krigdata, varmean, semivariogram, grid, grid_error, index0,
                        index1, interp, vardatamax, vardatamin):
    """
    :param krigdata:
    :param varmean_row:
    :param semivariogram:
    :param grid:
    :param grid_error:
    :param index0:
    :param index1:
    :param interp:
    :param vardatamax:
    :return:
    """
    if interp == "simple":
        try:
            # todo should i use varmean as I create a semivariogram for all field data?
            simple = krig.simple(krigdata, varmean, semivariogram)
        except Exception as e:
            print(e, end='')  # TODO did this because of the singular matrix error
        else:
            # TODO temporary patch to avoid outside range result
            if simple[0] < 0:
                grid[index0, index1] = vardatamin
                grid_error[index0, index1] = simple[1]
            elif simple[0] <= vardatamax * 2:
                grid[index0, index1] = simple[0]
                grid_error[index0, index1] = simple[1]


    else:  # ordinary
        try:
            ordinary = krig.ordinary(krigdata, semivariogram)
        except Exception as e:
            print(e, end='')  # TODO did this because of the singular matrix error
        else:
            # TODO this is a temporary patch to avoid outside range result!
            if ordinary[0] < 0:
                grid[index0, index1] = vardatamin #this will avoid vineyar rows with no values for high filtering
                grid_error[index0, index1] = ordinary[1]
            elif ordinary[0] <= vardatamax * 2:
                grid[index0, index1] = ordinary[0]
                grid_error[index0, index1] = ordinary[1]


def fillin_grids_colors(krigdata, varmean, semivariogram, grid, grid_error, index0,
                        index1, interp):
    """

    :param krigdata:
    :param varmean_row:
    :param semivariogram:
    :param grid:
    :param grid_error:
    :param index0:
    :param index1:
    :param interp:
    :return:
    """
    if interp == "simple":
        try:
            # todo should i use varmean as I create a semivariogram for all field data?
            simple = krig.simple(krigdata, varmean, semivariogram)
        except Exception as e:
            print(e, end='')  # TODO did this because of the singular matrix error
        else:
            # TODO temporary patch to avoid outside range result
            if simple[0] > 1:
                grid[index0, index1] = 1.0
                grid_error[index0, index1] = simple[1]
            elif simple[0] < 0:
                grid[index0, index1] = 0.0
                grid_error[index0, index1] = simple[1]
            else:
            #if 0 <= simple[0] <= 1:
                grid[index0, index1] = simple[0]
                grid_error[index0, index1] = simple[1]

    else:  # ordinary
        try:
            ordinary = krig.ordinary(krigdata, semivariogram)
        except Exception as e:
            print(e, end='')  # TODO did this because of the singular matrix error
        else:
            if ordinary[0] > 1:
                grid[index0, index1] = 1.0
                grid_error[index0, index1] = ordinary[1]
            elif ordinary[0] < 0:
                grid[index0, index1] = 0.0
                grid_error[index0, index1] = ordinary[1]
            # TODO this is a temporary patch to avoid outside range result!
            #if 0 <= ordinary[0] <= 1:
            else:
                grid[index0, index1] = ordinary[0]
                grid_error[index0, index1] = ordinary[1]


def dokriging(vardata, grid, grid_error, semivariogram, index, geotr, datasorce, interp, tree, radius):
    """

    :param vardata:
    :param grid:
    :param grid_error:
    :param semivariogram:
    :param index:
    :param geotr:
    :param datasorce:
    :param interp:
    :param tree: rtree
    :param radius: interpolation radius
    :return:
    """
    # get the number of pixels to iterate
    npixels = index[0].shape[0]

    for i in range(npixels):

        lonlat = gdal.ApplyGeoTransform(geotr, int(index[1][i]), int(index[0][i]))

        ###get the  nearest points inside the interpolation radius
        nearindexes = tree.query_ball_point(lonlat, radius, n_jobs=-1)
        # f.write(" "+str(len(nearindexes)))
        if not nearindexes:  # actually this shouldnt happen
            continue
        nearvardata = vardata[nearindexes]
        varmean = np.mean(nearvardata[:, 2])
        # this is the highest value in the nearby row data, this is used th threshod the output
        vardatamax = nearvardata[:, 2].max()
        vardatamin = nearvardata[:, 2].min()

        # if there is only 1 neighbour we take its value else we do kriging
        if len(nearindexes) == 1:
            # if len(idx) == 1:
            grid[index[0][i], index[1][i]] = nearvardata[0, 2]  # i did this otherwise kriging with 1 point will raise error
            grid_error[index[0][i], index[1][i]] = 0  # todo how to deal with the error?
        else:
            distance, idx = tree.query((lonlat[0], lonlat[1]), len(nearindexes))
            krigdata = np.hstack((nearvardata, np.expand_dims(distance, axis=0).T))

            # for now there is only global variogrm
            ############local variogram
            '''
            if variogram == "local":
                #"local calculate best variogram and kringing"

                # i set the hh to the pixel size
                hh = 1/area_multiplier
                # set lags from 0 to the interpolation radius
                #lags = range(0, int(radius), hh)
                lags = np.arange(0, int(radius), hh)
                gamma = vrc.semivar2(vardata, lags, hh)
                # covariance = vrc.covar(data, lags, hh)

                # choose the model with lowest rmse
                semivariograms = []
                semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.spherical))
                ###semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.linear))
                ###semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.gaussian))
                semivariograms.append(fit.fitsemivariogram(vardata, gamma, fit.exponential))
                rsmes = [rmse(gamma[0], i(gamma[0])) for i in semivariograms]
                semivariogram = semivariograms[rsmes.index(min(rsmes))]
            '''
            # if variogram == "byrow":
            #    pass

        if datasorce == 'counts':
            fillin_grids_counts(krigdata, varmean, semivariogram, grid, grid_error, index[0][i],
                                index[1][i], interp, vardatamax, vardatamin)
        else:
            fillin_grids_colors(krigdata, varmean, semivariogram, grid, grid_error, index[0][i],
                                index[1][i], interp)

    return grid, grid_error



def dokriging_byrow(dataset, row, step, field, grid, grid_error, semivariogram, random_filter, row_indexed, index, geotr, datasorce, interp, rowdirection):
    """
    Like the above dofriging, but this is for kriging by row
    :param dataset: pandas dataframe
    :param row:
    :param step:
    :param field:
    :param grid:
    :param grid_error:
    :param semivariogram:
    :param random_filter:
    :param row_indexed:
    :param index:
    :param geotr:
    :param datasorce: 'counts'  or 'colors'
    :param interp:
    :param rowdirection:
    :return:
    """

    # select one vineyard row
    rowdf = dataset[dataset["row"] == int(row)]

    # if random_filter we will set the radius to the max distance between points in a row
    if random_filter:
        max_point_distance = utils.max_point_distance(rowdf, "lon", "lat", "row",
                                                                    direction=rowdirection)
        radius = max_point_distance + 0.5
    else:
        radius = 0.5 * step

    # initialize the kdtree index for the current row
    xydata = rowdf[["lon", "lat"]].values
    tree = kdtree.cKDTree(xydata)

    vardata_row = rowdf[["lon", "lat", field]].values
    varmean_row = np.mean(vardata_row[:, 2])

    # set mask for the current row
    mask = row_indexed == int(row)
    index0_masked = index[0][mask]
    index1_masked = index[1][mask]

    # get the number of pixels to iterate
    npixels = index0_masked.shape[0]

    for i in range(npixels):

        lonlat = gdal.ApplyGeoTransform(geotr, int(index1_masked[i]), int(index0_masked[i]))

        ###get the  nearest points inside the interpolation radius
        nearindexes = tree.query_ball_point(lonlat, radius, n_jobs=-1)
        # f.write(" "+str(len(nearindexes)))
        if not nearindexes:  # actually this shouldnt happen
            continue
        nearvardata = vardata_row[nearindexes]
        varmean_row = np.mean(nearvardata[:, 2]) #########todo is this correct? ineed the mean of this subdata for simple kriging???
        # this is the ighest value in the nearby row data, this is used th threshod the output
        vardatamax = nearvardata[:, 2].max()
        vardatamin = nearvardata[:, 2].min()

        # if there is only 1 neighbour we take its value else we do kriging
        if len(nearindexes) == 1:
            # if len(idx) == 1:
            grid[index0_masked[i], index1_masked[i]] = nearvardata[0, 2]  # i did this otherwise kriging with 1 point will raise error
            grid_error[index0_masked[i], index1_masked[i]] = 0  # todo how to deal with the error?
        else:
            distance, idx = tree.query((lonlat[0], lonlat[1]), len(nearindexes))
            krigdata = np.hstack((nearvardata, np.expand_dims(distance, axis=0).T))

            # for now there is only global variogrm
            #if variogram == "local":
            #    pass
            #if variogram == "byrow":
            #    pass

            if datasorce == 'counts':

                fillin_grids_counts(krigdata, varmean_row, semivariogram, grid, grid_error, index0_masked[i],
                                    index1_masked[i], interp,vardatamax,vardatamin)
            else:
                fillin_grids_colors(krigdata, varmean_row, semivariogram, grid, grid_error, index0_masked[i],
                                    index1_masked[i], interp)

    return grid, grid_error




############################################   30 05 2017 ######################################

##add parameter that specify we want random filtered data for row  ( take first/last and then random in between the 2)
##add parameter to force inteerpolation for points in the same row only (for higher filters the interpolation radius will
##overlap with points belonging to adjacent rows

#so there are 8 possibilities for each interpolation method
# ordered filter, first/last, free interpolation
# ordered filter, no first/last, free interpolation
# ordered filter, first/last, interpolation by row only
# ordered filter, no first/last, interpolation by row only
# random filter, first/last, free interpolation
# random filter, no first/last, free interpolation
# random filter, first/last, interpolation by row only
# random filter, no first/last, interpolation by row only


#for random filter the interpolation radius is calculated with the max distance between points on the same row


#the interpolation radius is used for the semivariogram lags max distance


def berrycolor_workflow_kriging(steps=[30], colorgrades = ['a','b','c','d'], counts = ['raw','visible'], interp="simple",
                                random_filter=False, first_last=False, force_interpolation_by_row=False, nodata=-1.0,
                                variogram = "global",rowdirection="x", epsg="32611", rounddecimals=4, forcenugget=True,
                                id = "adca0a053062b45da8a20e0a0d4b4fe54", skip=True):

    """
    Do kriging interpolation on the colorgrade data and output rasters and statistics file

    :param steps: the filtering steps to use
    :param colorgrades:  colorgrades to consider    possibles ['a','b','c','d']
    :param counts: count fields to consider   possibles ['raw','visible']
    :param interp: "simple" or "ordinary"
    :param random_filter: true to take random points (e,g step2 takes 50% random points for each row, false use ordered filter
    (e.g. step2 takes 1 point every 2 points)
    :param first_last: true to always have the first and last point of a vineyard row
    :param force_interpolation_by_row: True if i want to be sure the points belong to the current row only and no
                                        points from other rows are used during interpolation
    :param nodata:
    :param variogram: "global","local","byrow"  (only global is now supported, byrow is like local but only for points from
                                                the current row)

                                                        variogram
                                            global      byrow       local
    force_interpolation_by_row    true      X           X           X
                                  false     X                       X
    :param rowdirection:    'x' ,'y'
    :param epsg: coordinate system code
    :param rounddecimals: number of decimals in the resulting iterpolation grids
     this is used to round the decimals in the result grids, this should limit overflows when calculating the statistics
    :param forcenugget True to force the distance zero for the interpolated semivariogram to have the nugget value
    :param id: used for debugging purposes
    :param skip: used for debugging purposes
    :return:
    """

    area_multiplier = 2  # 2 to double the resulution and halve the pixel size to 0.5 meters

    if variogram not in ["global", "byrow", "local"]:
        raise ValueError("Variogram should be 'global' or 'byrow' or 'local'")

    if variogram == "byrow":
        raise NotImplementedError("variogram by row is not implemented")
    if variogram == "local":
        raise NotImplementedError("local variogram cannot be used at present")

    if variogram == 'byrow' and not force_interpolation_by_row:
        raise ValueError("Variogram=='byrow' cannot be used with force_interpolation_by_row==False")


    if interp not in ["simple", "ordinary"]:
        raise ValueError("Interpolation should be 'simple' or 'ordinary'")



    #jsonpath = os.path.join(os.path.dirname(__file__), '../..', "pysdss/experiments/interpolation/")
    jsonpath = os.path.join(os.path.dirname(__file__), '../../../..', "pysdss/experiments/interpolation/")
    jsonpath = os.path.normpath(jsonpath)



    folder = "/vagrant/code/pysdss/data/output/text/"
    experimentfolder = "/vagrant/code/pysdss/experiments/interpolation/"

    # create a folder to store the output
    id = utils.create_uniqueid()
    if not os.path.exists(folder + id):
        os.mkdir(folder + id)
    folder = folder + id + "/"


    outinfo(folder+"info.json",
            "simple kriging" if interp=="simple" else "ordinary kriging",
            "random" if random_filter else "ordered",
            "onlyrows" if force_interpolation_by_row else "free",
            first_last,
            steps,
            colorgrades,
            counts,
            variogram
            )

    # create a directory to store settings
    # if not os.path.exists(experimentfolder + str(id)):
    #   os.mkdir(experimentfolder + str(id))
    # dowload_data_test(id, folder)

    # dowload the data #TODO: create function to read from the database
    df = pd.read_csv("/vagrant/code/pysdss/data/input/2016_06_17.csv",
                     usecols=["%Dataset_id", " Row", " Raw_fruit_count",
                              " Visible_fruit_per_m", " Latitude", " Longitude", "Harvestable_A", "Harvestable_B",
                              "Harvestable_C", "Harvestable_D"])
    df.columns = ["id", "row", "raw_fruit_count", "visible_fruit_per_m", "lat", "lon", "a", "b", "c", "d"]

    # here should be necessary some mapping to get the column names


    # filter data (and save to disk .
    # 1/03/17  we don't filter
    # keep, discard = filter.filter_byvalue(df, 0, ">", colname="d")
    keep = df

    # save filtered data to disk after reprojection
    # reproject latitude, longitude #WGS 84 / UTM zone 11N
    west, north = utils2.reproject_coordinates(keep, "epsg:"+epsg, "lon", "lat", False)
    filter.set_field_value(keep, west, "lon")
    filter.set_field_value(keep, north, "lat")
    filter.set_field_value(keep, 0, "keepa")
    filter.set_field_value(keep, 0, "keepb")
    filter.set_field_value(keep, 0, "keepc")
    filter.set_field_value(keep, 0, "keepd")


    ##############search and remove duplicates
    #keep = utils.remove_duplicates(keep, ["lat", "lon"], operation="mean")
    ########################################

    keep.to_csv(folder + id + "_keep.csv", index=False)

    '''west,north =utils.reproject_coordinates(discard, "epsg:32611", "lon", "lat", False)
    filter.set_field_value(discard, west,"lon")
    filter.set_field_value(discard, north,"lat")
    filter.set_field_value(discard, 0, "keepd")
    discard.to_csv(folder + id +"_discard.csv",index=False)'''


    # create the virtualdataset header for the datafile

    with open(jsonpath + "/vrtmodel.txt") as f:
        xml = f.read()

    # output the 4 virtual files for the 4 columns
    # for clr in ["a"]:
    for clr in ["a", "b", "c", "d"]:
        data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
                "elevation": clr}
        utils2.make_vrt(xml, data, folder + id + "_keep" + clr + ".vrt")

    # output the virtual files for raw fruit count and visible fruit count
    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "raw_fruit_count"}
    utils2.make_vrt(xml, data, folder + id + "_keep_rawfruit.vrt")

    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "visible_fruit_per_m"}
    utils2.make_vrt(xml, data, folder + id + "_keep_visiblefruit.vrt")

    # 5 define an interpolation grid (the user can choose an existing grid or draw on screen)
    # here i am doing manually for testing, i just take the bounding box

    minx = math.floor(keep['lon'].min())
    maxx = math.ceil(keep['lon'].max())
    miny = math.floor(keep['lat'].min())
    maxy = math.ceil(keep['lat'].max())
    delty = maxy - miny  # height
    deltx = maxx - minx  # width



    # log a running process to the database
    # save grid metadata and interpolation metadata to the database historic tables
    # xecute gridding with gdal_array (set nodata value to zero or other value)

    path = jsonpath + "/average.json"
    params = utils.json_io(path, direction="in")

    params["-txe"] = str(minx) + " " + str(maxx)
    params["-tye"] = str(miny) + " " + str(maxy)
    params["-a"]["nodata"] = str(nodata)
    # pixel is 0.5 meters
    params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)

    # for clr in ["a"]:
    for clr in colorgrades:

        params["src_datasource"] = folder + id + "_keep" + clr + ".vrt"
        params["dst_filename"] = folder + id + "_" + clr + ".tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        # call gdal_grid
        print("Getting the 'truth' raster for color " + clr)
        out, err = utils.run_tool(text)
        print("output" + out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    ##rasterize raw fruit and visible fruit

    if 'raw' in counts:
        params["src_datasource"] = folder + id + "_keep_rawfruit.vrt"
        params["dst_filename"] = folder + id + "_rawfruit.tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        # call gdal_grid
        print("Getting the 'truth' raster for raw fruit count")
        out, err = utils.run_tool(text)
        print("output" + out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    if 'visible' in counts:
        params["src_datasource"] = folder + id + "_keep_visiblefruit.vrt"
        params["dst_filename"] = folder + id + "_visiblefruit.tiff"

        # build gdal_grid request
        text = grd.build_gdal_grid_string(params, outtext=False)
        print(text)
        text = ["gdal_grid"] + text
        print(text)

        # call gdal_grid
        print("Getting the 'truth' raster for visible fruit per m")
        out, err = utils.run_tool(text)
        print("output" + out)
        if err:
            print("error:" + err)
            raise Exception("gdal grid failed")

    # 9 add result url to the database table
    # 10 log completed process (maybe should add the location of the result


    # 11 define the indexed array for the images, I used the same interpolation method and parameters therefore the index
    #  is the same

    # for clr in ["a"]:
    for clr in colorgrades:
        d = gdal.Open(folder + id + "_" + clr + ".tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                              nodata=-1)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id + "_indexedarray_" + clr, (index, r_indexed, properties))
        d = None


    if 'raw' in counts:
        # define the indexed array for the raw and visible fruit
        d = gdal.Open(folder + id + "_rawfruit.tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                              nodata=-1)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id + "_indexedarray_rawfruit", (index, r_indexed, properties))
        d = None


    if 'visible' in counts:
        d = gdal.Open(folder + id + "_visiblefruit.tiff")
        index, r_indexed, properties = gdalIO.raster_dataset_to_indexed_numpy(d, id, maxbands=1, bandLocation="byrow",
                                                                              nodata=-1)
        print("saving indexed array to disk")
        fileIO.save_object(folder + id + "_indexedarray_visiblefruit", (index, r_indexed, properties))
        d = None

    ######## rasterize rows with nearest neighbour

    data = {"layername": id + "_keep", "fullname": folder + id + "_keep.csv", "easting": "lon", "northing": "lat",
            "elevation": "row"}
    utils2.make_vrt(xml, data, folder + id + "_keep_row.vrt")
    '''newxml = xml.format(**data)
    f = open(folder + id+"_keep_row.vrt", "w")
    f.write(newxml)
    f.close()'''

    path = jsonpath + "/nearest.json"
    params = utils.json_io(path, direction="in")

    params["-txe"] = str(minx) + " " + str(maxx)
    params["-tye"] = str(miny) + " " + str(maxy)
    params["-outsize"] = str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
    params["-a"]["nodata"] = str(nodata)
    params["src_datasource"] = folder + id + "_keep_row.vrt"
    params["dst_filename"] = folder + id + "_row.tiff"

    # build gdal_grid request
    text = grd.build_gdal_grid_string(params, outtext=False)
    print(text)
    text = ["gdal_grid"] + text
    print(text)

    # call gdal_grid
    print("Getting the row raster")
    out, err = utils.run_tool(text)
    print("output" + out)
    if err:
        print("error:" + err)
        raise Exception("gdal grid failed")

    # upload to numpy
    d = gdal.Open(folder + id + "_row.tiff")
    band = d.GetRasterBand(1)
    # apply index
    row_indexed = gdalIO.apply_index_to_single_band(band,index)  # this is the last index from the fruit count, all the indexes are the same
    d = None

    # collect statistics for the comparisons
    stats = {}
    # mean=[]
    # std=[]

    #######################################################

    ####### get the indexed array for raw count
    if 'raw' in counts:
        index_raw, r_indexed_raw, properties_raw = fileIO.load_object(folder + id + "_indexedarray_rawfruit")

    ####### get the indexed array for visible count
    if 'visible' in counts:
        index_visible, r_indexed_visible, properties_visible = fileIO.load_object(folder + id + "_indexedarray_visiblefruit")

    ###########################
    # upload to numpy
    d = gdal.Open(folder + id + "_row.tiff")
    band = d.GetRasterBand(1)
    # apply index
    row_indexed = gdalIO.apply_index_to_single_band(band,index)  # this is the last index from the fruit count, all the indexes are the same


    geotr = d.GetGeoTransform()    # get the grotransfor to be used later during kriging, looks like gdal_grid set the origin in the bottom left!!!

    d = None


    if force_interpolation_by_row:

        ## make temporary directory to store the interpolated vineyard rows
        #################################################
        tempfolder = folder + "/temprasters/"
        os.mkdir(tempfolder)

        # define new rasters increasing the data filtering and interpolating
        for step in steps:

            # filter data by step
            k, d = filter.filter_bystep(keep, step=step, rand=random_filter, first_last=first_last,rowname="row")
            k.to_csv(folder + id + "_keep_step" + str(step) + ".csv", index=False)
            d.to_csv(folder + id + "_discard_step" + str(step) + ".csv", index=False)

            # find unique row numbers
            rows = k["row"].unique()

            # add dictionary for this step
            stats[step] = {}
            for clr in colorgrades:
                stats[step][clr] = {}
                for row in rows:
                    stats[step][clr][row] = {}
                    stats[step][clr][row]['truth'] = []
                    stats[step][clr][row]['interpolated'] = []


            #calculate max distance between points to, to make things easier I am not doing it row by row
            if random_filter: max_point_distance_all = utils.max_point_distance(k, "lon", "lat", "row", direction=rowdirection)
            #extract the data to be used to build the variogram
            vardata = k[["lon", "lat", "visible_fruit_per_m"]].values #todo what about raw?
            varmean = np.mean( vardata[:, 2]) #todo is this necessary here?
            # here we calculate the global variogram for visible_fruit_per_m
            if variogram == "global":
                hh = 1 / area_multiplier
                maxlagdistance = max_point_distance_all + 0.5 if random_filter else 0.5 * step
                outfile = folder + id + "_globalsemivariogram_" + "visible" + "_step" + str(step)
                semivariogram = get_variogram(vardata, hh, maxlagdistance, outfile, forcenugget=forcenugget)


            for clr in counts:

                for row in rows:

                    # initialize grid as numpy array, filled in with nodata value (-1)
                    # initialize grid as numpy array, filled in with nodata value (-1)
                    grid = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)
                    grid_error = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)

                    # execute kriging
                    #todo should be done for raw too?
                    grid, grid_error = dokriging_byrow(k, row, step, "visible_fruit_per_m", grid, grid_error, semivariogram, random_filter,
                                                       row_indexed, index, geotr, "counts", interp, rowdirection)

                    # delete values less than zero and round the values
                    grid[grid < 0] = nodata
                    np.round(grid, rounddecimals, out=grid)

                    # save grids to raster
                    xsize = deltx * area_multiplier
                    ysize = delty * area_multiplier
                    gridpath = tempfolder + id + "_" + clr + "_step" + str(step) +  "_row" + str(row) + "_" + interp + variogram
                    grid_errorpath = tempfolder + id + "_" + clr + "_step" + str(step) +  "_row" + str(row) + "_" + interp + variogram + "_error"

                    savegrids(grid, gridpath, grid_error, grid_errorpath, xsize, ysize, geotr, nodata, epsg)


            # extract statistics for raw and visible
            for clr in colorgrades:


                for count in counts:

                    print("calculate statistics for "+count+" berrie step" + str(step) + "color" + clr)
                    for row in rows:

                        '''
                        # upload to numpy
                        d = gdal.Open(tempfolder + id + "_raw_step" + str(step) + "_row"+str(row) + ".tiff")
                        band = d.GetRasterBand(1)
                        # apply index
                        new_r_indexed_raw = gdalIO.apply_index_to_single_band(band,
                                                                              index)  # this is the last index from the fruit count
                        d = None
                        '''

                        d = gdal.Open(tempfolder + id + "_" + count + "_step" + str(step) +  "_row" + str(row) + "_" + interp + variogram + ".tiff")
                        band = d.GetRasterBand(1)
                        # apply index
                        new_r_indexed_visible = gdalIO.apply_index_to_single_band(band,
                                                                                  index)  # this is the last index from the fruit count


                        d = None

                        if new_r_indexed_visible.min() == nodata:
                            warnings.warn(
                                "indexed data visible berry per meters has nodata values, the current implementation will"
                                " count this pixels as nozero values", RuntimeWarning)
                            new_r_indexed_visible[new_r_indexed_visible == nodata] = 'nan'


                        #update structure with partial data
                        # get a mask for the current row
                        mask = row_indexed == row
                        # statistics for current row
                        # average, std, totalarea,total nonzero area, total zeroarea, total raw fruit count,
                        # average raw fruit count, std raw fruit count ,total visible fruitxm, average visible fruitXm, std visible fruitXm

                        stats[step][clr][row]['truth'] = [None, None, None, None, None,
                                              #r_indexed_raw[0, :][mask].sum(),
                                              #r_indexed_raw[0, :][mask].mean(),
                                              #r_indexed_raw[0, :][mask].std(),
                                              r_indexed_visible[0, :][mask].sum(),
                                              r_indexed_visible[0, :][mask].mean(),
                                              r_indexed_visible[0, :][mask].std()]

                        #stats[step][clr][row]['interpolated'] = [None, None, None, None, None,
                                                     #new_r_indexed_raw[mask].sum(),
                                                     #new_r_indexed_raw[mask].mean(),
                                                     #new_r_indexed_raw[mask].std(),
                        #                             new_r_indexed_visible[mask].sum(),
                        #                             new_r_indexed_visible[mask].mean(),
                        #                             new_r_indexed_visible[mask].std()]

                        stats[step][clr][row]['interpolated'] = [None, None, None, None, None,
                                                            # np.nansum(new_r_indexed_raw[mask]),
                                                            # np.nanmean(new_r_indexed_raw[mask]),
                                                            # np.nanstd(new_r_indexed_raw[mask]),
                                                            np.nansum(new_r_indexed_visible[mask]),
                                                            np.nanmean(new_r_indexed_visible[mask]),
                                                            np.nanstd(new_r_indexed_visible[mask])]


            for clr in colorgrades:

                ###calculate semivariogram , we calculate from all the points
                # extract the data to be used to build the variogram
                vardata = k[["lon", "lat", clr.lower()]].values
                varmean = np.mean(vardata[:, 2]) #todo is this necessary here?
                # here we calculate the global variogram for visible_fruit_per_m
                if variogram == "global":
                    hh = 1 / area_multiplier
                    maxlagdistance = max_point_distance_all + 0.5 if random_filter else 0.5 * step
                    outfile = folder + id + "_globalsemivariogram_" + clr.lower() + "_step" + str(step)
                    semivariogram = get_variogram(vardata, hh, maxlagdistance, outfile, forcenugget=forcenugget)


                #for each row if random calculate the max distance for the current row
                #execute interpolation for jus the row points
                #calculate statistics for the row and fill in the output data structure

                for row in rows:

                    # initialize grid as numpy array, filled in with nodata value (-1)
                    # initialize grid as numpy array, filled in with nodata value (-1)
                    grid = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)
                    grid_error = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)

                    #do kriging
                    grid, grid_error = dokriging_byrow(k, row, step, clr.lower(), grid, grid_error, semivariogram, random_filter,
                                                       row_indexed, index, geotr, "colors", interp, rowdirection)

                    # round the values to avoid overflows when calculating statistics
                    np.round(grid, rounddecimals, out=grid)

                    # save grids to raster
                    xsize = deltx * area_multiplier
                    ysize = delty * area_multiplier
                    gridpath = tempfolder + id + "_" + clr + "_step" + str(step) + "_row" + str(row) + "_" + interp + variogram
                    grid_errorpath = tempfolder + id + "_" + clr + "_step" + str(step) + "_row" + str(row) + "_" + interp + variogram + "_error"

                    savegrids(grid, gridpath, grid_error, grid_errorpath, xsize, ysize, geotr, nodata, epsg)


                    # upload to numpy
                    d = gdal.Open(tempfolder + id + "_" + clr + "_step" + str(step) + "_row" + str(row) + "_" + interp + variogram + ".tiff")
                    band = d.GetRasterBand(1)
                    # apply index
                    new_r_indexed = gdalIO.apply_index_to_single_band(band,index)  # this is the last index from the fruit count

                    d = None

                    if new_r_indexed.min() == nodata:
                        warnings.warn(
                            "indexed data for colorgrade " + clr + " has nodata values, the current implementation will"
                                                                   " count this pixels as nozero values",RuntimeWarning)
                        new_r_indexed[new_r_indexed == nodata] = 'nan'  # careful nan is float


                    # calculate average, std, total area (here a pixel is 0.25 m2,
                    # totla area for pixel > 0, total area for pixels with 0 value

                    ####### get the indexed array for this colour grade

                    # d = gdal.Open(folder + id + "_" + clr + ".tiff")
                    # band = d.GetRasterBand(1)
                    # r_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count
                    # d = None
                    index, r_indexed, properties = fileIO.load_object(folder + id + "_indexedarray_" + clr)

                    # update structure with partial data
                    area = math.pow(1 / area_multiplier, 2)
                    # get a mask for the current row
                    mask = row_indexed == row
                    #fill in with data
                    partial_truth_data = stats[step][clr][row]['truth']
                    partial_truth_data[:5] = [r_indexed[0, :][mask].mean(), r_indexed[0, :][mask].std(),
                    r_indexed[0, :][mask].shape[0] * area,
                    np.count_nonzero(r_indexed[0, :][mask]) * area,
                    r_indexed[0, :][mask][r_indexed[0, :][mask] == 0].shape[0] * area]

                    partial_interpolated_data = stats[step][clr][row]['interpolated']

                    '''partial_interpolated_data[:5] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                    new_r_indexed[mask].shape[0] * area,
                    np.count_nonzero(new_r_indexed[mask]) * area,
                    new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area]'''

                    partial_interpolated_data[:5] = [np.nanmean(new_r_indexed[mask]), np.nanstd(new_r_indexed[mask]),
                                           new_r_indexed[mask].shape[0] * area,
                                           np.count_nonzero(new_r_indexed[mask]) * area,
                                           new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area]


    else: ################## no force interpolation by row

        # define new rasters increasing the data filtering and interpolating
        for step in steps:
            # add dictionary for this step
            stats[step] = {}

            #start = time.time()

            # filter data by step
            k, d = filter.filter_bystep(keep, step=step, rand=random_filter, first_last=first_last,rowname="row")
            # if random_filter we will set the radius to the max distance between points in a row
            if random_filter: max_point_distance = utils.max_point_distance(k, "lon", "lat", "row", direction=rowdirection)

            '''
            i already projected before
            west, north = utils.reproject_coordinates(k, "epsg:32611", "lon", "lat", False)
            filter.set_field_value(k, west, "lon")
            filter.set_field_value(k, north, "lat")
            filter.set_field_value(k, 0, "keepa")
            filter.set_field_value(k, 0, "keepb")
            filter.set_field_value(k, 0, "keepc")
            filter.set_field_value(k, 0, "keepd")'''
            k.to_csv(folder + id + "_keep_step" + str(step) + ".csv", index=False)

            '''west, north = utils.reproject_coordinates(d, "epsg:32611", "lon", "lat", False)
            filter.set_field_value(d, west, "lon")
            filter.set_field_value(d, north, "lat")
            filter.set_field_value(d, 0, "keepa")
            filter.set_field_value(d, 0, "keepb")
            filter.set_field_value(d, 0, "keepc")
            filter.set_field_value(d, 0, "keepd")'''
            d.to_csv(folder + id + "_discard_step" + str(step) + ".csv", index=False)

            if random_filter: radius = max_point_distance + 0.5
            else: radius = 0.5 * step

            ''' grid parameters to build the numpy array
            str(minx) + " " + str(maxx)
            str(miny) + " " + str(maxy)
            str(deltx * area_multiplier) + " " + str(delty * area_multiplier)
            '''

            #########################

            # initialize the global kdtree index (we'll use this also for the colors
            xydata = k[["lon", "lat"]].values
            tree = kdtree.cKDTree(xydata)

            # extract the data to be used to build the variogram
            vardata = k[["lon", "lat", "visible_fruit_per_m"]].values
            varmean = np.mean(vardata[:, 2]) #todo is this useful?

            #now= time.time()
            #print( "start variogram " + str(now - start *1000))

            if variogram == "global":
                hh = 1 / area_multiplier
                outfile = folder + id + "_globalsemivariogram_" + "visible" + "_step" + str(step)
                semivariogram = get_variogram(vardata, hh, radius, outfile, forcenugget=forcenugget)

            #now1= time.time()
            #print( "finish variogram " + str(now1 - now *1000))

            # initialize grid as numpy array, filled in with nodata value (-1)
            grid = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)
            grid_error = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)

            grid, grid_error = dokriging(vardata, grid, grid_error, semivariogram, index, geotr, "counts", interp, tree,radius)


            #now= time.time()
            #print(" kriging over " + str(now - now1 *1000))

            # delete values less than zero and round the values
            grid[grid < 0] = nodata
            np.round(grid, rounddecimals, out=grid)

            # save grids to raster
            xsize = deltx * area_multiplier
            ysize = delty * area_multiplier
            gridpath = folder + id + "_visible_step" + str(step) + "_" + interp + variogram
            grid_errorpath = folder + id + "_visible_step" + str(step) + "_" + interp + variogram + "_error"
            savegrids(grid, gridpath, grid_error, grid_errorpath, xsize, ysize, geotr, nodata, epsg)


            '''
            # upload to numpy
            d = gdal.Open(folder + id + "_raw_step" + str(step) + ".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed_raw = gdalIO.apply_index_to_single_band(band,
                                                                  index)  # this is the last index from the fruit count
            d = None
            '''

            d = gdal.Open(folder + id + "_visible_step" + str(step) + "_" + interp + variogram + ".tiff")
            band = d.GetRasterBand(1)
            # apply index
            new_r_indexed_visible = gdalIO.apply_index_to_single_band(band,
                                                                      index)  # this is the last index from the fruit count
            d = None

            if new_r_indexed_visible.min() == nodata:
                warnings.warn(
                    "indexed data visible berry per meters has nodata values, the current implementation will"
                                                           " count this pixels as nozero values", RuntimeWarning)
                new_r_indexed_visible[new_r_indexed_visible == nodata] = 'nan'


            # for clr in ["a"]:
            for clr in colorgrades:

                ############################### colors
                vardata = k[["lon", "lat", clr.lower()]].values
                varmean = np.mean(vardata[:, 2]) #todo is this useful?

                #now1 = time.time()
                #print(now1 - now * 1000)

                if variogram == "global":
                    hh = 1 / area_multiplier
                    outfile = folder + id + "_globalsemivariogram_" + clr.lower() + "_step" + str(step)
                    semivariogram = get_variogram(vardata, hh, radius, outfile, forcenugget=forcenugget)

                #now = time.time()
                #print(now - now1 * 1000)

                # initialize grid as numpy array, filled in with nodata value (-1)
                grid = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)
                grid_error = np.ones((delty * area_multiplier, deltx * area_multiplier)) * (nodata)

                grid, grid_error = dokriging(vardata, grid, grid_error, semivariogram, index, geotr, "colors", interp,tree, radius)

                #now1 = time.time()
                #print(now1 - now * 1000)

                # round the values to avoid overflows when calculating statistics
                np.round(grid, rounddecimals, out=grid)

                # save grids to raster
                xsize = deltx * area_multiplier
                ysize = delty * area_multiplier
                gridpath = folder + id + "_" + clr + "_step" + str(step) + "_" + interp + variogram
                grid_errorpath = folder + id + "_" + clr + "_step" + str(step) + "_" + interp + variogram + "_error"
                savegrids(grid, gridpath, grid_error, grid_errorpath, xsize, ysize, geotr, nodata, epsg)

                # upload to numpy
                d = gdal.Open(folder + id + "_"+clr+"_step" + str(step) + "_" + interp + variogram + ".tiff")
                band = d.GetRasterBand(1)
                # apply index
                new_r_indexed = gdalIO.apply_index_to_single_band(band,
                                                                  index)  # this is the last index from the fruit count
                d = None

                # check if all pixels have a value, otherwise assign nan to nodata value
                if new_r_indexed.min() == nodata:
                    warnings.warn(
                        "indexed data for colorgrade " + clr + " has nodata values, the current implementation will"
                                                               " count this pixels as nozero values", RuntimeWarning)
                    new_r_indexed[new_r_indexed == nodata] = 'nan'  # careful nan is float


                # calculate average, std, total area (here a pixel is 0.25 m2,
                # total area for pixel > 0, total area for pixels with 0 value

                ####### get the indexed array for this colour grade

                # d = gdal.Open(folder + id + "_" + clr + ".tiff")
                # band = d.GetRasterBand(1)
                # r_indexed = gdalIO.apply_index_to_single_band(band, index) # this is the last index from the fruit count
                # d = None
                index, r_indexed, properties = fileIO.load_object(folder + id + "_indexedarray_" + clr)

                stats[step][clr] = {}

                for i in np.unique(row_indexed):  # get the row numbers

                    area = math.pow(1 / area_multiplier, 2)

                    # add dictionary for this row
                    stats[step][clr][i] = {}

                    # get a mask for the current row
                    mask = row_indexed == i
                    # statistics for current row

                    # average, std, totalarea,total nonzero area, total zeroarea, total raw fruit count,
                    # average raw fruit count, std raw fruit count ,total visible fruitxm, average visible fruitXm, std visible fruitXm

                    # r_indexed is 2d , while new_r_indexed and mask are 1d
                    stats[step][clr][i]['truth'] = [r_indexed[0, :][mask].mean(), r_indexed[0, :][mask].std(),
                                                    r_indexed[0, :][mask].shape[0] * area,
                                                    np.count_nonzero(r_indexed[0, :][mask]) * area,
                                                    r_indexed[0, :][mask][r_indexed[0, :][mask] == 0].shape[0] * area,
                                                    #r_indexed_raw[0, :][mask].sum(),
                                                    #r_indexed_raw[0, :][mask].mean(),
                                                    #r_indexed_raw[0, :][mask].std(),
                                                    r_indexed_visible[0, :][mask].sum(),
                                                    r_indexed_visible[0, :][mask].mean(),
                                                    r_indexed_visible[0, :][mask].std()]


                    ''' 
                    stats[step][clr][i]['interpolated'] = [new_r_indexed[mask].mean(), new_r_indexed[mask].std(),
                                                           new_r_indexed[mask].shape[0] * area,
                                                           np.count_nonzero(new_r_indexed[mask]) * area,
                                                           new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area,
                                                           #new_r_indexed_raw[mask].sum(),
                                                           #new_r_indexed_raw[mask].mean(),
                                                           #new_r_indexed_raw[mask].std(),
                                                           new_r_indexed_visible[mask].sum(),
                                                           new_r_indexed_visible[mask].mean(),
                                                           new_r_indexed_visible[mask].std()]
                    '''

                    stats[step][clr][i]['interpolated'] = [ np.nanmean(new_r_indexed[mask]),  np.nanstd(new_r_indexed[mask]),
                                                            new_r_indexed[mask].shape[0] * area,
                                                            np.count_nonzero(new_r_indexed[mask]) * area,
                                                            new_r_indexed[mask][new_r_indexed[mask] == 0].shape[0] * area,
                                                            # np.nansum(new_r_indexed_raw[mask]),
                                                            # np.nanmean(new_r_indexed_raw[mask]),
                                                            # np.nanstd(new_r_indexed_raw[mask]),
                                                            np.nansum(new_r_indexed_visible[mask]),
                                                            np.nanmean(new_r_indexed_visible[mask]),
                                                            np.nanstd(new_r_indexed_visible[mask])]


                    # compare = r_indexed[0,:][mask] - new_r_indexed[mask]
                    # stats[clr][step][i]['comparison'] = [compare.mean(),compare.std()]
                    # pass
                    # calculate average and variance for this surface
                    # mean.append(compare.mean())
                    # std.append(compare.std())

    # print(mean)
    # print(std)
    return id, stats
    # {2: {"a": {1: {'truth': {}, 'interpolated': {}}}}}




####OUTPUT STATISTICS#################



def test_interpret_result(stats, id, folder, colmodel=None):
    '''
    Interpret the dictionary with the results and save a table, statistics are per row and files are per color grade

    output file names are like <colorgrade>_output.csv


    table columns are:
    row|avg|std|area|nozero_area|zero_area|raw_count|raw_avg|raw_std|visible_count|visible_avg|visible_std
    step2avg|step2std|step2area|step2nozero_area|step2zero_area|step2raw_count|step2raw_avg|step2raw_std|step2visible_count|step2visible_avg|step2visible_std
    |step3.....

    :param stats:
        dictionary in the form
        {step:{"colorgrade": { row : {'truth':{},'interpolated':{} }} } }
    :param rsme: True to calculate rootmeansquare error
    :return: None
    '''

    # get the number of steps and their value
    nsteps = len(stats)
    print(nsteps)
    b = [i for i in stats]
    b.sort()
    print(b)

    # define the column names and convert to string
    if not colmodel:
        colmodel = ["avg", "std", "area","nozero_area", "zero_area", "raw_count","raw_avg","raw_std","visible_count","visible_avg","visible_std"]

    NCOL = len(colmodel)
    cols =  ["row"] + colmodel
    for s in b:
        cols += ["step"+str(s)+"_"+i for i in colmodel]
    print(cols)
    print(len(cols))


    # calculate the number of columns for the result
    ncolumns = NCOL * (nsteps + 1) + 1
    print(ncolumns)

    # get the name of the colorgrades
    c = [i for i in stats[b[0]]]
    c.sort()
    print(c)

    #get the number of rows
    nrows = len(stats[b[0]][c[0]])
    print(nrows)

    #step=1
    for j in c: #for each colorgrade we want to save in a separate file

        # create a numpy array to store result
        output = np.zeros((nrows, ncolumns), dtype=np.float32)

        # first step, color grade j; output both the truth and the first step
        f = stats[b[0]][ j ]

        r = 0 if 0 in f else 1 #if the column starts at 1 we need to know (the array index start at zero)
        for row in f:  # iterate rows, each one has truth and interpolated
            row = int(row)
            nrow = row-r #correct the index for the rows

            output[nrow, 0] = row

            truth = f[row]["truth"]
            for i in range(NCOL):
                output[nrow, i+1] = truth[i]

            interp = f[row]["interpolated"]
            for i in range(NCOL):
                output[nrow, i + 1 + NCOL] = interp[i]

        # now we can output the remaining steps
        if nsteps >1:
            for n in range(1, nsteps):
                f = stats[b[n]][j]
                for row in f:
                    nrow = row - r #correct the index for the rows
                    interp = f[row]["interpolated"]
                    for i in range(NCOL):
                        output[nrow, i + 1 + (NCOL*2) + NCOL*(n-1)] = interp[i] # 1row + 5truth + 5firststep + other steps

        outcols = ','.join(cols)
        #save array to disk as a csv for this color grade
        np.savetxt(folder + id +"_"+j+"_output.csv",output , fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='', comments='')


def test_compare_raster_totruth(point, id, folder, epsg=32611, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10],
                                removeduplicates=True, clean=None,
                                colorgrade=None, counts=None, suffix="" ):
    """
    comparing original point values with the underlying pixel values and save comparison to disk
    use this when the workflow was using force_interpolation_by_row=False

    output file name is like _ERRORSTATISTICS.csv

    :param point: this is the id +"_keep.csv" file
    :param id: operation id
    :param folder: the folder with point
    :param epsg: epsg code for spatial reference
    :param steps: the steps to consider
    :param removeduplicates: true to get average value for duplicate points
    :param clean: chage first value with the second [[-1],[0]] , useful to clean nodata values
    :param tempfolder: name of the folder storing the rasters by row
    :return:
    """

    # step = [2, 3, 4,10]

    if not colorgrade:
        colorgrade = ["a", "b", "c", "d"]
    if not counts:
        counts = ["raw", "visible"]

    # 1 open  csv

    # id = "a16b2a828b9174d678e76be46619eb329"
    # folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    # point = folder + id + "_keep.csv"


    if removeduplicates:  # TODO check that remove dulicates will not alter the row number
        # remove duplicates and save to newfile, set name for point file
        df = utils.remove_duplicates(point, ["lat", "lon"], operation="mean")
        point = folder + id + "_keep_nodup.csv"
        df.to_csv(point)

    df = pd.read_csv(point, usecols=["id", "row", "a", "b", "c", "d", "raw_fruit_count", "visible_fruit_per_m"])

    # 2 extract


    # define a list of files
    rasters = [folder + id + "_" + y + "_step" + str(x) + suffix + ".tiff" for x in steps for y in colorgrade]
    rasters2 = [folder + id + "_" + y + "_step" + str(x) + suffix + ".tiff" for x in steps for y in counts]


    spatialRef = osr.SpatialReference()
    spatialRef.ImportFromEPSG(epsg)
    srs = spatialRef.ExportToProj4()
    vrtdict = {"layername": os.path.basename(point).split(".")[0], "fullname": point, "easting": "lon",
               "northing": "lat", "srs": srs}

    c = []
    for r in rasters:
        stat = utils2.compare_csvpoint_cell(point, r, vrtdict, overwrite=False)
        c.append(stat)

    arr = np.transpose(np.asarray(c))
    newdf = pd.DataFrame(arr, columns=["step" + str(x) + y for x in steps for y in colorgrade])

    c = []
    for r in rasters2:
        stat = utils2.compare_csvpoint_cell(point, r, vrtdict, overwrite=False)
        c.append(stat)

    arr = np.transpose(np.asarray(c))
    newdf2 = pd.DataFrame(arr, columns=["step" + str(x) + y for x in steps for y in counts])

    outdf = pd.concat([df, newdf, newdf2], axis=1)

    if clean:  # clean the dataframe, useful to hide nodata values in the result  e.g convert -1 with 0
        utils.clean_dataframe(outdf, clean[0], clean[1])  # TODO check the use of  DataFrame.replace

    outdf.to_csv(folder + id + "_ERRORSTATISTICS.csv")


def test_compare_raster_totruth_byrow(point, id, folder, epsg=32611, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10],
                                      removeduplicates=True,
                                      clean=None, tempfolder="/temprasters/", colorgrade=None, counts=None, suffix="" ):
    """
    comparing original point (with no filter) values with the underlying pixel values and save comparison to disk
    use this when the workflow was using force_interpolation_by_row=True

    output file name is like _ERRORSTATISTICS.csv


    :param point: this is the id +"_keep.csv" file
    :param id: operation id
    :param folder: the folder with point
    :param epsg: epsg code for spatial reference
    :param steps: the steps to consider
    :param removeduplicates: true to get average value for duplicate points
    :param clean: chage first value with the second [[-1],[0]] , useful to clean nodata values
    :param tempfolder: name of the folder storing the rasters by row
    :return:
    """

    # step = [2, 3, 4,10]
    if not colorgrade:
        colorgrade = ["a", "b", "c", "d"]
    if not counts:
        counts = ["raw", "visible"]

    # 1 open  csv
    # id = "a16b2a828b9174d678e76be46619eb329"
    # folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    # point = folder + id + "_keep.csv"

    if removeduplicates:  # TODO check that remove dulicates will not alter the row number
        # remove duplicates and save to newfile, set name for point file
        df = utils.remove_duplicates(point, ["lat", "lon"], operation="mean")
        point = folder + id + "_keep_nodup.csv"
        df.to_csv(point, index=False)

    df = pd.read_csv(point,
                     usecols=["id", "row", "a", "b", "c", "d", "raw_fruit_count", "visible_fruit_per_m", "lat", "lon"])
    # find unique row numbers
    rows = sorted(df["row"].unique())
    # filter by row and save to tempfolder
    for row in rows:
        row = int(row)
        rowdf = df[df["row"] == row]
        # save csv file to disk
        rowdf.to_csv(folder + "/temprasters/" + id + "_keep_row" + str(row) + ".csv", index=False)
    df.drop(["lat", "lon"], axis=1, inplace=True)  # we dont want in the final result

    # 2 extract
    # define a dictionary of files
    rasters = {}
    rasters2 = {}
    for row in rows:
        row = int(row)
        # this will be in order a,b,c,d   and raw,visible
        rasters[row] = [folder + tempfolder + id + "_" + y + "_step" + str(x) + "_row" + str(row) + suffix + ".tiff" for x in
                        steps for y in colorgrade]
        rasters2[row] = [folder + tempfolder + id + "_" + y + "_step" + str(x) + "_row" + str(row) + suffix + ".tiff" for x in
                         steps for y in counts]

    # define a gdal spatial reference object
    spatialRef = osr.SpatialReference()
    spatialRef.ImportFromEPSG(epsg)
    srs = spatialRef.ExportToProj4()

    c = []
    for row in rows:
        row = int(row)
        orig_points = folder + "/temprasters/" + id + "_keep_row" + str(row) + ".csv"
        vrtdict = {"layername": os.path.basename(orig_points).split(".")[0], "fullname": orig_points, "easting": "lon",
                   "northing": "lat", "srs": srs}
        raster_row = []
        for r in rasters[row]:
            stat = utils2.compare_csvpoint_cell(orig_points, r, vrtdict, overwrite=True)
            raster_row.append(stat)  # this will be a linear collecion of values step2a, step2b,....
        c.append(np.transpose(np.asarray(raster_row)))

    # we stack vertically all the rows and convert to dataframe
    arr_colour = np.vstack(c)
    newdf = pd.DataFrame(arr_colour, columns=["step" + str(x) + y for x in steps for y in colorgrade])

    c = []
    for row in rows:
        row = int(row)
        orig_points = folder + "/temprasters/" + id + "_keep_row" + str(row) + ".csv"
        vrtdict = {"layername": os.path.basename(orig_points).split(".")[0], "fullname": orig_points, "easting": "lon",
                   "northing": "lat", "srs": srs}
        raster_row = []
        for r in rasters2[row]:
            # overwrite false because we already have the shapefiles
            stat = utils2.compare_csvpoint_cell(orig_points, r, vrtdict, overwrite=False)
            raster_row.append(stat)
        c.append(np.transpose(np.asarray(raster_row)))

    arr_count = np.vstack(c)
    newdf2 = pd.DataFrame(arr_count, columns=["step" + str(x) + y for x in steps for y in counts])

    # stack the 3 dataframe horizontally
    outdf = pd.concat([df, newdf, newdf2], axis=1)

    if clean:  # clean the dataframe, useful to hide nodata values in the result  e.g convert -1 with 0
        utils.clean_dataframe(outdf, clean[0], clean[1])  # TODO check the use of  DataFrame.replace

    outdf.to_csv(folder + id + "_ERRORSTATISTICS.csv", index=False)


def estimate_berries_byrow(folder,id, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10]):

    """ Output comparison between original points (with no filter) and interploated points for average, std, estimated number per row

    data comes from _ERRORSTATISTICS.csv

    will output file _count_stats_byrow.csv

    :param id:
    :param folder:
    :param steps:
    :return:
    """

    left_fields = ["row","length","visible_m_avg_original","visible_m_std_original","estimated_berries_original"]
    right_fields = ["visible_m_avg", "delta_avg", "visible_std", "delta_std", "estimated_berries", "delta_estimates"]


    # calculate row length
    lengths = utils.row_length(folder + id + "_keep.csv", "lon", "lat", "row", conversion_factor=1)
    # initialize output array
    out = np.zeros((lengths.shape[0], 5 + (len(steps)+1)*6))
    out[:,0:2] = lengths

    # open  _ERRORSTATISTICS and calculate average and std per row for the visible count
    path = folder + id + "_ERRORSTATISTICS.csv"
    df = pd.read_csv(path, usecols=['row', 'visible_fruit_per_m'])

    # open one of the files with the statistics for the rasterized data
    path = folder + id + "_d_output.csv"
    df2 = pd.read_csv(path)

    #gey statistics for original points and rasterized points
    rows = lengths[:,0]
    for i,row in enumerate(rows):
        # select one vineyard row
        fdf = df[df['row'] == row]

        out[i, 2] = fdf['visible_fruit_per_m'].mean()
        out[i, 3] = fdf['visible_fruit_per_m'].std()
        out[i, 4] = out[i, 1] * out[i, 2]

        out[i, 5] = df2["visible_avg"].values[i]
        out[i, 6] = np.absolute(out[i, 5] - out[i, 2])
        out[i, 7] = df2["visible_std"].values[i]
        out[i, 8] = np.absolute(out[i, 7] - out[i, 3])
        out[i, 9] = out[i, 1] * out[i, 5]
        out[i, 10] = np.absolute(out[i, 9] -  out[i, 4])

    # get statistics for interpolated filtered points
    idxstart = 11
    for s in steps:
        for i, row in enumerate(rows):
            out[i, idxstart] = df2[ "step"+str(s)+"_visible_avg"].values[i]
            out[i, idxstart+1] = np.absolute(out[i, idxstart] - out[i, 2])
            out[i, idxstart+2] = df2["step"+str(s)+"_visible_std"].values[i]
            out[i, idxstart+3] = np.absolute(out[i, idxstart+2] - out[i, 3])
            out[i, idxstart+4] = out[i, 1] * out[i, idxstart]
            out[i, idxstart+5] = np.absolute(out[i, idxstart+4] - out[i, 4])
        idxstart += len(right_fields)

    # save array to disk
    cols = left_fields + right_fields
    cols += ["step"+str(x)+"_"+ i  for x in steps for i in right_fields ]
    outcols = ','.join(cols)
    np.savetxt(folder + id + "_count_stats_byrow.csv", out, fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='',
               comments='')



def estimate_colors_byrow(folder,id, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10], colorgrades = ["a", "b", "c", "d"]):

    """ Output comparison between original points (with no filter) and interploated points for average,

    data comes from _ERRORSTATISTICS.csv

    will output file _color_stats_byrow.csv

    :param folder:
    :param id:
    :param steps:
    :params colorgrades:
    :return: None
    """

    left_fields = ["row"]
    right_fields = []
    for c in colorgrades:
        right_fields.append(c+"_avg")
    left_fields = left_fields + right_fields

    #"_rsme_row.csv contains row and x_avg, the average from the original point data
    path = folder + id + "_rsme_row.csv"
    df = pd.read_csv(path, usecols= left_fields+right_fields)

    # initialize output array
    out = np.zeros((df.shape[0], len(left_fields) + len(right_fields)*2*len(steps)))
    out[:, : len(left_fields)] = df.values[:,:len(left_fields)]


    startindex = len(left_fields)
    for c in colorgrades:
        # x_output contains  "stepn_avg"
        path = folder + id + "_"+c+"_output.csv"
        cls = [ "step"+ str(s) +"_avg" for s in steps]
        df2 = pd.read_csv(path, usecols= cls)

        for s in steps:

            out[:,startindex] = df2["step"+ str(s) +"_avg"].values
            out[:, startindex+1] = np.absolute(out[:,startindex] - df[ c +"_avg"].values)

            startindex+=2

    # save array to disk
    for c in colorgrades:
        for s in steps:
            left_fields +=[ "step"+ str(s)+"_"+ c +"_avg","step"+ str(s)+"_"+ c +"_delta_avg"]

    outcols = ','.join(left_fields)
    np.savetxt(folder + id + "_color_stats_byrow.csv", out, fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='',
               comments='')

########CALCULATE RMSE#######################


def calculate_rmse_row_0(folder, id, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10], colorgrades = ["a", "b", "c", "d"]):
    """
    rmse based on comparison with the rasterized points
    """
    #open a random csv as pandas
    path = folder + id + "_" + random.choice(colorgrades) + "_output.csv"
    file = pd.read_csv(path)

    outtext = ""

    # calculate rmse for visible count (this is the same for all colour grade files)
    for e,c in enumerate(steps):
        outtext += "step"+str(c)+"_visible_avg\t" + str(utils.rmse( file["step"+str(c)+"_visible_avg"].values, file["visible_avg"].values)) + "\n"

    for e,c in enumerate(colorgrades):
        #for each step
        path = folder + id + "_" + c + "_output.csv"
        file = pd.read_csv(path)
        for j,x in enumerate(steps):

            outtext += "step" + str(x) + "_avg\t"+ str(utils.rmse(file["step" + str(x) + "_avg"].values, file["avg"].values)) + "\n"

    with open(folder + id + "_rmse_row.csv","w") as f:
        f.write(outtext)


def calculate_rmse_row(folder, id, steps=[2, 3, 4, 5, 6, 7, 8, 9, 10], colorgrades = ["a", "b", "c", "d"]):
    """
    rmse based on comparison with the original points (this is preferred to calculate_rmse_row_0)
    rmse are calculated for each row, the number of items are the points of the row
    the starting file is the ERRORSTATISTICS.csv file which contains values for the original points and the underlying pixels
    :param folder: 
    :param id: 
    :param steps:
    :param colorgrades: 
    :return: 
    """

    rowname = "row"
    stepcount = len(steps)
    colorcount = len(colorgrades)

    path = folder + id + "_ERRORSTATISTICS.csv"
    df = pd.read_csv(path)

    #get name unique rows and initialize an empty array for the output
    rows = df[rowname].unique()
    arr = np.zeros((len(rows), 2 + stepcount+ colorcount + stepcount*colorcount))

    #iterate the row
    for i,row in enumerate(rows):
        # select one vineyard row
        fdf = df[df[rowname] == row]

        #set value for row and average visible_fruit_per_m for the current row
        arr[i,0] = row
        arr[i, 1] = fdf['visible_fruit_per_m'].mean()

        # calculate rmse for visible count
        for e, c in enumerate(steps):
            arr[i,2 + e] = utils.rmse(fdf["step" + str(c) + "visible"].values, fdf["visible_fruit_per_m"].values)

        #set average value for colorgrades for the current row
        for e,c in enumerate(colorgrades):
            idxstart = 2 + stepcount
            arr[i, idxstart + e] = fdf[c].mean()

        # calculate root mean square error for colorgrades
        for e,c in enumerate(colorgrades):
            idxstart = 2 + stepcount + colorcount + (e * stepcount)
            for j, x in enumerate(steps):
                arr[i, idxstart + j] = utils.rmse(fdf["step" + str(x) + c].values, fdf[c].values)

    #set xolumn names and save as a csv file
    cols = ['row', 'visible_avg']
    cols += ["visible_step" + str(s) +"_rmse" for s in steps]
    cols +=  [i + "_avg" for i in colorgrades]
    cols += ["step" + str(i) + "_" + c + "_rmse" for c in colorgrades for i in steps]
    outcols = ','.join(cols)
    np.savetxt(folder + id + "_rsme_row.csv",arr , fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='', comments='')


def calculate_rmse(folder,id,steps=[2,3,4,5,6,7,8,9,10]):
    """
    Calculate the root mean square error for the visible per meter average and the estimated berries . The differences
    are calculated per number of vineyard rows and the number of items are the rows
    :param folder: 
    :param id: 
    :param nsteps: 
    :return: 
    """

    #create pandas dataframe
    path = folder + id + "_count_stats_byrow.csv"
    df = pd.read_csv(path)

    # number of vineyard rows
    nrows = df.shape[0]

    # store the output field named
    outcols_0 = []
    outcols_1 = []

    #cretae empty numpy array
    arr = np.zeros((1, len(steps)*2))

    #iterate the steps and fill in the array
    for i,s in enumerate(steps):

        #rmse for visible_avg, add to the output array, append the column name
        x =  math.sqrt((np.sum(np.power(df["step"+str(s)+"_delta_avg"].values,2))/nrows))
        arr[0, i] = x
        outcols_0.append("step"+str(s)+"_delta_avg_rmse")
        #rmse for estimated berries, add to the output array, append the column name
        y = math.sqrt((np.sum(np.power(df["step"+str(s)+"_delta_estimates"].values, 2)) / nrows))
        arr[0, i+len(steps)] = y
        outcols_1.append("step"+str(s)+"_delta_estimates_rmse")

    outcols_0 += outcols_1
    outcols = ','.join(outcols_0)
    #export array to disk
    np.savetxt(folder + id + "_rsme.csv",arr , fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='', comments='')


def calculate_rmse_color(folder,id,steps=[2,3,4,5,6,7,8,9,10],colorgrades = ["a", "b", "c", "d"]):
    """
    Calculate the root mean square error for the colorgrade delta average  . The differences
    are calculated per number of vineyard rows and the number of items are the rows

    output  _n_rmse.csv for each colorgrade
    :param folder:
    :param id:
    :param nsteps:
    :return:
    """

    #create pandas dataframe
    path = folder + id + "_color_stats_byrow.csv"
    df = pd.read_csv(path)

    # number of vineyard rows
    nrows = df.shape[0]



    #cretae empty numpy array
    arr = np.zeros((1, len(steps)))


    for c in colorgrades:

        # store the output field named
        outcols_0 = []
        #iterate the steps and fill in the array
        for i,s in enumerate(steps):

            #rmse for visible_avg, add to the output array, append the column name
            x =  math.sqrt((np.sum(np.power(df["step"+str(s)+"_"+c+"_delta_avg"].values,2))/nrows))
            arr[0, i] = x
            outcols_0.append("step"+str(s)+"_delta_avg_rmse")

        outcols = ','.join(outcols_0)
        #export array to disk
        np.savetxt(folder + id + "_" + c + "_rsme.csv",arr , fmt='%.3f', delimiter=',', newline='\n', header=outcols, footer='', comments='')



#############CHARTING#####################


def chart_rmse(folder,id, useperc=True):
    """Output charts for rmse delta_avg number of berries   and  delta_estimates number of berries
    :param folder:
    :param id:
    :param useperc: True to output percentages, else output the step
    :return: None
    """

    #open pandas dataframe
    path = folder + id + "_rsme.csv"
    df = pd.read_csv(path)

    ncols = df.shape[1]
    #the categorical labels for x axis (step2,step3,step4,....)
    x_labels = [i.split("_")[0] for i in df.columns][:int(len(df.columns) / 2)]

    #x_label_num =  [ int(n[4:]) for n in x_labels].sort()
    #x_labels_all =  ["step"+str(i) for i in range(2,x_label_num[-1]+1)]

    #convert to percentage
    if useperc:
        x_labels = [str(int(1/int(i[4:])*100))+"%" for i in x_labels]

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    #chart ["stepx_delta_avg_rmse"]
    output_file(folder + "/charts/delta_avg_rmse.html",title="delta_avg_rmse")
    f = figure(x_range=x_labels, title="Delta RMSE for average number of berries per vineyard row")
    f.line(x_labels , df.values[0,0:int(ncols/2)], line_width=2)
    save(f)

    #chart ("stepx_delta_estimates_rmse")
    output_file(folder + "/charts/delta_estimates_rmse.html",title="delta_estimates_rmse")
    f = figure(x_range=x_labels, title="Delta RMSE for total number of berries per vineyard row")
    f.line(x_labels , df.values[0,int(ncols/2):], line_width=2)
    save(f)


def chart_rmse_color(folder,id, colorgrades = ["a", "b", "c", "d"], useperc=True):
    """Output charts for rmse delta_avg number of berries   and  delta_estimates number of berries
    :param folder:
    :param id:
    :param useperc: True to output percentages, else output the step
    :return: None
    """


    for c in colorgrades:

        #open pandas dataframe
        path = folder + id + "_"+c+"_rsme.csv"
        df = pd.read_csv(path)

        ncols = df.shape[1]
        #the categorical labels for x axis (step2,step3,step4,....)
        x_labels = [i.split("_")[0] for i in df.columns]

        #x_label_num =  [ int(n[4:]) for n in x_labels].sort()
        #x_labels_all =  ["step"+str(i) for i in range(2,x_label_num[-1]+1)]

        #convert to percentage
        if useperc:
            x_labels = [str(int(1/int(i[4:])*100))+"%" for i in x_labels]

        #make new chart directory if it does not exist
        if not os.path.exists(folder +"/charts"):
            os.mkdir(folder + "/charts")

        #chart ["stepx_delta_avg_rmse"]
        output_file(folder + "/charts/"+c+"_delta_avg_rmse.html",title=c+"_delta_avg_rmse")
        f = figure(x_range=x_labels, title="Delta RMSE for average colourgrade "+c.upper()+" per vineyard row")
        f.line(x_labels , df.values[0,:ncols], line_width=2)
        save(f)



def compare_chart_rmse(folder,id,folder1,id1, outpath, colorgrade="", suffix="", outstats=None,exportimage=True, useperc=True):
    """
    Compare the rmse for 2 datasets (calculate_rmse() should be used beforehand to get the rmse)
    :param folder: first file folder
    :param id: first file id
    :param folder1: second file folder
    :param id1: second file id
    :param outpath: where to output the charts
    :param:suffix additional text to add to the file name
    :param outstats a file object for review infos
    :param useperc: use "n%" in the chart otherwise use "stepXX"
    :return:
    """

    if exportimage:
        outchart = outpath+"/charts"
        if not os.path.exists(outchart): os.mkdir(outchart)

    ################# first dataset
    #open pandas dataframe
    if colorgrade: path = folder + "/" + id + "_" + colorgrade + "_rsme.csv"
    else:path = folder + "/" + id + "_rsme.csv"
    df = pd.read_csv(path)

    ncols = df.shape[1]
    #the categorical labels for x axis (step2,step3,step4,....)
    if colorgrade:  x_labels = [int(i.split("_")[0][4:]) for i in df.columns]
    else: x_labels = [int(i.split("_")[0][4:]) for i in df.columns][:int(len(df.columns) / 2)]

    #x_label_num =  [ int(n[4:]) for n in x_labels].sort()
    #x_labels_all =  ["step"+str(i) for i in range(2,x_label_num[-1]+1)]

    #convert to percentage
    if useperc:
        x_labels = [int((1/i)*100) for i in x_labels]

    ################### second dataset
    # open pandas dataframe
    if colorgrade: path = folder1 + "/" + id1 + "_" + colorgrade + "_rsme.csv"
    else: path = folder1 + "/" + id1 + "_rsme.csv"
    df1 = pd.read_csv(path)

    ncols1 = df1.shape[1]
    # the categorical labels for x axis (step2,step3,step4,....)
    if colorgrade:  x_labels1 = [int(i.split("_")[0][4:]) for i in df1.columns]
    else: x_labels1 = [int(i.split("_")[0][4:]) for i in df1.columns][:int(len(df1.columns) / 2)]

    # x_label_num =  [ int(n[4:]) for n in x_labels].sort()
    # x_labels_all =  ["step"+str(i) for i in range(2,x_label_num[-1]+1)]

    # convert to percentage
    if useperc:
        x_labels1 = [int((1/i)*100) for i in x_labels1]
    #########################

    ######prepare the labels
    if useperc:
        xlabels = sorted(list(set(x_labels + x_labels1)), reverse=True)
        xlabels = [str(i) + "%" for i in xlabels]

        x_labels = [str(i) + "%" for i in x_labels]
        x_labels1 = [str(i) + "%" for i in x_labels1]

    else:
        xlabels = sorted(list(set(x_labels + x_labels1)))
        xlabels = ["step"+str(i) for i in xlabels]

        x_labels = ["step"+str(i) for i in x_labels]
        x_labels1 = ["step"+str(i) for i in x_labels1]


    #get infos
    info = utils.json_io(folder+"/info.json", direction="in")
    info1 = utils.json_io(folder1+"/info.json", direction="in")


    ############################
    #chart ["stepx_delta_avg_rmse"]

    dupl1 = ""
    dupl2 = ""
    legend1 = info["name"]
    legend2 = info1["name"]
    if info["name"] == info1["name"]:
        legend1 += "1"
        legend2 += "2"
        dupl1 = "1"
        dupl2 = "2"

    #these are info to add to output chart name
    addendum = ("T" if info["filter_type"]=="random" else "F") + ("T" if info["interpolation"] =="onlyrows" else "F") + ("T" if info["first_last"]==True else "F")
    addendum1 = ("T" if info1["filter_type"]=="random" else "F") + ("T" if info1["interpolation"] =="onlyrows" else "F") + ("T" if info1["first_last"]==True else "F")

    if colorgrade:
        fname= "/"+colorgrade+"_delta_avg_rmse_comparison"+suffix+ "_" + info["name"].split()[0] + addendum + "_" + info1["name"].split()[0] + addendum1
        output_file(outpath+fname+".html",title=colorgrade+"_delta_avg_rmse_comarison")
        f = figure(x_range=xlabels, title="Delta RMSE for colourgrade "+colorgrade.upper()+" per vineyard row")
        f.line(x_labels , df.values[0,:ncols], line_width=2, color="blue",legend=legend1)
        f.line(x_labels1 , df1.values[0,:ncols1], line_width=2,color="green",legend=legend2)
    else:
        fname= "/delta_avg_rmse_comparison"+suffix+ "_" + info["name"].split()[0] + addendum + "_" + info1["name"].split()[0] + addendum1
        output_file(outpath+fname+".html",title="delta_avg_rmse_comarison")
        f = figure(x_range=xlabels, title="Delta RMSE for average number of berries per vineyard row")
        f.line(x_labels , df.values[0,0:int(ncols/2)], line_width=2, color="blue",legend=legend1)
        f.line(x_labels1 , df1.values[0,0:int(ncols1/2)], line_width=2,color="green",legend=legend2)


    # making multiline labels /n and <br/> do not work
    first_label = Label(x=50, y=f.plot_height-50, x_units='screen', y_units='screen',
                       text=info["name"].upper() + dupl1+ " -> Filter:" + info["filter_type"] +" - Interpolation:" + info["interpolation"]+" - First/Last:" + str(info["first_last"]),
                        render_mode='canvas',background_fill_color='white', background_fill_alpha=1.0,text_font_size="7.5pt")

    second_label = Label(x=50, y=f.plot_height-70, x_units='screen', y_units='screen',
                         text=info1["name"].upper() + dupl2 + " -> Filter:" + info1["filter_type"] + " - Interpolation:" + info1["interpolation"]+" - First/Last:" + str(info1["first_last"]),
                        render_mode='canvas',background_fill_color='white', background_fill_alpha=1.0, text_font_size="7.5pt")

    f.add_layout(first_label)
    f.add_layout(second_label)
    f.legend.label_text_font_size = "7pt"
    f.legend.location = "bottom_right"
    save(f)

    if exportimage: export_png(f, filename=outchart + fname+".png")


    #chart ("stepx_delta_estimates_rmse")
    if colorgrade:
        pass
    else:
        fname = "/delta_estimates_rmse_comparison"+suffix+ "_" + info["name"].split()[0] + addendum + "_" + info1["name"].split()[0] + addendum1
        output_file(outpath+fname+".html",title="delta_estimates_rmse_comparison")
        f = figure(x_range=xlabels, title="Delta RMSE for total number of berries per vineyard row")
        f.line(x_labels, df.values[0,int(ncols/2):], line_width=2, color="blue",legend=legend1)
        f.line(x_labels1, df1.values[0,int(ncols1/2):], line_width=2,color="green", legend=legend2)

        first_label = Label(x=60, y= f.plot_height-50, x_units='screen', y_units='screen',
                           text=info["name"].upper() + dupl1 + " -> Filter:" + info["filter_type"] +" - Interpolation:" + info["interpolation"]+" - First/Last:" + str(info["first_last"]),
                            render_mode='canvas',background_fill_color='white', background_fill_alpha=1.0,text_font_size="7.5pt")

        second_label = Label(x=60, y= f.plot_height-70, x_units='screen', y_units='screen',
                             text=info1["name"].upper() + dupl2 + " -> Filter:" + info1["filter_type"] + " - Interpolation:" + info1["interpolation"]+" - First/Last:" + str(info1["first_last"]),
                            render_mode='canvas',background_fill_color='white', background_fill_alpha=1.0,text_font_size="7.5pt")
        f.add_layout(first_label)
        f.add_layout(second_label)
        f.legend.label_text_font_size = "7pt"
        f.legend.location = "bottom_right"
        save(f)

        if exportimage: export_png(f, filename=outchart + fname + ".png")

    #write info to the output file
    if outstats:
        outstats.write(suffix+ "," + id + "," + info["name"].split()[0] + "," +addendum + "," +  id1 + "," + info1["name"].split()[0] + "," + addendum1+"\n")


def compare_rmse_couples(paths, outpath, colorgrade=""):
    """
    Compare RMSE for each couple and output a camparison chart plus a text file with infos

    the overview.csv file structure is:
        id	id1	file1	file1_info	id2	file2	file2_info
    0	a00bca573171a4ff2a8ab70d58438dec0	simple	TFT	a01b2eada44b8426391d94c5249a025ab	moving	TFF
    1	a00bca573171a4ff2a8ab70d58438dec0	simple	TFT	a066f6d0e9a0d49d2999c3714c5f9ed34	inverse	FTF

    :param paths: a list of paths to the folders
    :param outpath: the folder to store the charts
    :return:
    """

    #id = "a40b9fc7181184bf0b9835f830a677513"
    #folder = "/vagrant/code/pysdss/data/output/text/analysis0806_steps2to100/" + id + "/"
    #id1 = "a7aa8d3e709ce42a4ba17ebee02daf66f"
    #folder1 = "/vagrant/code/pysdss/data/output/text/analysys1906_krigingsteps5to100/" + id1 + "/"
    #outpath = "/vagrant/code/pysdss/data/output/text/"

    #open a file to write infos about the comparison charts
    outstats = open(outpath+"/overview.csv","w")
    outstats.write("id,id1,file1,file1_info,id2,file2,file2_info\n")

    #make all the couple combinations for file names
    combinations = itertools.combinations(paths, 2)
    for i,path in enumerate(combinations):
        b = os.path.basename(path[0])
        id = b if b.strip() != '' else path[0].split('/')[-2]
        b = os.path.basename(path[0])
        id1 = b if b.strip() != '' else path[1].split('/')[-2]

        if colorgrade:compare_chart_rmse(path[0], id, path[1], id1, outpath,colorgrade=colorgrade, suffix= str(i), outstats=outstats, useperc=True)
        else:compare_chart_rmse(path[0], id, path[1], id1, outpath, suffix= str(i), outstats=outstats, useperc=True)
    outstats.close()



def publish_rsme_comparisons(path,outpath,typ="",keys=[], returnlist=False):
    """
    Find pictures for avg or estimates rmse, filter by keys, and publish as tables in a Word filw (.docx)

    -pass one key to find all the images with that key
    -pass2 keys for filtering both subcomponents of the image
        -the same key  "simple","simple"     "TFT","TFT"
        -different keys  "simpleTFT", "TTT"
        -one is subkey   "simpleTFT", "TFT"


    image names have this structure:
    - delta_avg_rmse_comparison0_simpleTFT_movingTFF.png
    - delta_estimates_rmse_comparison427_inverseTFF_inverseTTF.png


    :param path: path to folder with images only
    :param outpath: the folder to output the .docx file
    :param typ: "avg"  or  "estimates", this will be the first filter
    :param keys: list with one ot 2 query keys, this will be the second filter
    :param returnlist: True to return a list of the filtered files
    :return: a list of filtered file names if returnlist==True
    """

    #find all images
    files = [d for d in os.listdir(path) if not os.path.isdir(os.path.join(path, d))]
    # filter "averages" or "estimates" charts
    files = [d for d in files if typ in d]
    #filter by keys
    if len(keys) < 1: raise ValueError("pass 2 keys")
    if len(keys)==1: #only 1 key, just find everything with that key
        files = [d for d in files if (keys[0] in d)]
    else:
        if keys[0] == keys[1]: #same key, be sure the key is available ywice in the file
            files = [d for d in files if (len([m for m in re.finditer(keys[0], d)]) > 1)]
        else:
            #if keys[0] in keys[1] and len(keys[0]) == 3 : #we have 1 subkey then be sure this is available twice in the file
            if keys[0] in keys[1]:
                files = [d for d in files if (keys[1] in d and len([m for m in re.finditer(keys[0], d)]) > 1)]
            #elif keys[1] in keys[0] and len(keys[1]) == 3:#we have 1 subkey then be sure this is available twice in the file
            elif keys[1] in keys[0]:
                files = [d for d in files if (keys[0] in d and len([m for m in re.finditer(keys[1], d)]) > 1)]
            else: #2 different keys
                files = [d for d in files if (keys[0] in d and keys[1] in d)]  # filter "averages" or "estimates" charts

    #make the .docx document , add a table , and put images inside the table
    document = Document()
    table = document.add_table(cols=2, rows=math.ceil(len(files)/2)) #calculate the number of cells for a table with 2 rows
    i=0
    for row in table.rows:
        for cell in row.cells:
            if i<len(files):
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(path+"/"+files[i], width=Cm(7.5))
                i+=1

    #output name different if 1 or 2 keys
    if len(keys)==1:
        document.save(outpath + "/" + typ + "_" + keys[0] + ".docx")
    else:
        document.save(outpath+"/" + typ + "_"+ keys[0] + "_" + keys[1] +".docx")

    #print(files)
    if returnlist:
        return files


def grab_image_fromhtml(path, filename):
    """
    Grab a bokeh chart from an html page with phantojs(must be installed on th system)
    :param path: full path to the html page
    :param filename: full path to png output file
    :return: None
    """

    driver = webdriver.PhantomJS()
    driver.get("file:///" + path)

    ## resize for PhantomJS compat
    driver.execute_script("document.body.style.width = '100%';")

    #get a screenshot of the page and then resize to the bokeh chart boundaries
    png = driver.get_screenshot_as_png()
    bounding_rect_script = "return document.getElementsByClassName('bk-root')[0].children[0].getBoundingClientRect()"
    b_rect = driver.execute_script(bounding_rect_script)

    driver.quit()

    #crop and save the image
    image = Image.open(io.BytesIO(png))
    cropped_image = image.crop((b_rect["left"], b_rect["top"], b_rect["right"], b_rect["bottom"]))
    cropped_image.save(filename)


def export_estimated_berries_tables(rootdir, dirs,outpath):
    """
    This function will output a word document with tables, table will store the charts with the comparison of the
    estimated berries for each vineyard row at a specific step
    the charts are grabbed from the  "/charts/estimatedberries/estimated_stepXX.html" files
    :param rootdir: the root directory with the interpolation directories
    :param dirs: a list of interpolation directories
    :param outpath: the path to output the .docx file
    :return:
    """

    document = Document() #word document

    fulldirs = [rootdir + dir + "/charts/estimatedberries/" for dir in dirs]
    for k,path in enumerate(fulldirs):
        # find all html pages
        print("finding html files")
        files = [d for d in os.listdir(path) if not os.path.isdir(os.path.join(path, d))]

        steps = [int(d.split("estimated_step")[1].split(".")[0]) for d in files]
        steps.sort()

        print("make a table")
        #make the .docx document , add a table , and put images inside the table
        document.add_paragraph(dirs[k])
        table = document.add_table(cols=2, rows=math.ceil(len(files)/2)) #calculate the number of cells for a table with 2 rows
        i=0
        for row in table.rows:
            for cell in row.cells:
                if i<len(files):
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    #save a tmp image as png and add to the table
                    grab_image_fromhtml(path +"/estimated_step"+str(steps[i])+".html", outpath+"/tmp.png")
                    run.add_picture(outpath+"/tmp.png", width=Cm(7.5))
                    i+=1
        document.add_page_break()

    document.save(outpath + "/estimated_berries_tables.docx")


def export_estimated_colors_tables(rootdir, dirs,outpath, colorgrades=["a","b","c","d"]):
    """
    This function will output a word document with tables, table will store the charts with the comparison of the
    estimated colorgrades for each vineyard row at a specific step
    the charts are grabbed from the  "/charts/estimatedberries/estimated_stepXX.html" files
    :param rootdir: the root directory with the interpolation directories
    :param dirs: a list of interpolation directories
    :param outpath: the path to output the .docx file
    :return:
    """

    document = Document() #word document

    for c in colorgrades:

        fulldirs = [rootdir + dir + "/charts/colorgrades/colors" + c.upper() + "_comparisons" for dir in dirs]
        for k,path in enumerate(fulldirs):
            # find all html pages
            print("finding html files")
            files = [d for d in os.listdir(path) if not os.path.isdir(os.path.join(path, d))]

            steps = [int(d.split("estimated_" + c.upper() + "_step")[1].split(".")[0]) for d in files]
            steps.sort()

            print("make a table")
            #make the .docx document , add a table , and put images inside the table
            document.add_paragraph(dirs[k])
            table = document.add_table(cols=2, rows=math.ceil(len(files)/2)) #calculate the number of cells for a table with 2 rows
            i=0
            for row in table.rows:
                for cell in row.cells:
                    if i<len(files):
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        #save a tmp image as png and add to the table
                        grab_image_fromhtml(path +"/estimated_" + c.upper() + "_step"+str(steps[i])+".html", outpath+"/tmp.png")
                        run.add_picture(outpath+"/tmp.png", width=Cm(7.5))
                        i+=1
            document.add_page_break()

    document.save(outpath + "/estimated_colors_tables.docx")



def chart_estimated_berries(folder,id, steps=[], useperc=True):
    """Output charts of the estimated berries original data vs filtered data
    :param folder: folder path
    :param id: operation id
    :param steps: list with steps to consider
    :param useperc: True to label steps as percentages
    :return: None
    """

    path = folder + id + "_count_stats_byrow.csv"

    cols = ['row', 'estimated_berries'] + ["step"+str(s)+"_estimated_berries" for s in steps]

    df = pd.read_csv(path, usecols=cols)

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries"):
        os.mkdir(folder + "/charts/estimatedberries")

    steps_perc = [int(1 / int(s) * 100) for s in steps] if useperc else steps

    for i,s in enumerate(steps):
        output_file(folder +"/charts/estimatedberries/estimated_step"+str(s)+".html", title = "estimated_step"+str(s))
        f = figure(title='Estimated berries: original VS '+str(steps_perc[i])+"%") if useperc else \
            figure(title='Estimated berries: original VS step'+str(s))
        f.line(df["row"], df["estimated_berries"], line_color="green", line_width=1, legend="Original")
        f.line(df["row"], df["step"+str(s)+"_estimated_berries"], line_color="red", line_width=1, legend=str(steps_perc[i])+"%" if useperc else "Step"+str(s))
        f.legend.location = "top_left"
        f.yaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.axis_label = "vineyard rows"
        f.yaxis.axis_label = "number of berries"
        save(f)


def chart_estimated_colors(folder,id, steps=[], colorgrades=["a","b","c","d"], useperc=True):
    """Output charts of the estimated average colorgrade, original data vs filtered data
    :param folder: folder path
    :param id: operation id
    :param steps: list with steps to consider
    :param colorgrades: colorgraes to consider
    :param useperc: True to label steps as percentages
    :return: None
    """

    path = folder + id + "_color_stats_byrow.csv"

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    # make new chart directory if it does not exist
    if not os.path.exists(folder + "/charts/colorgrades"):
        os.mkdir(folder + "/charts/colorgrades")

    cols = ["row"]
    for c in colorgrades:
        cols.append(c+"_avg")
    for c in colorgrades:
        cols = cols + ["step"+str(s)+ "_" + c + "_avg" for s in steps]

    df = pd.read_csv(path, usecols=cols)


    for c in colorgrades:

        #make new chart directory if it does not exist
        if not os.path.exists(folder +"/charts/colorgrades/colors" + c.upper() + "_comparisons"):
            os.mkdir(folder + "/charts/colorgrades/colors" + c.upper() + "_comparisons")

        steps_perc = [int(1 / int(s) * 100) for s in steps] if useperc else steps

        for i,s in enumerate(steps):
            output_file(folder +"/charts/colorgrades/colors" + c.upper() + "_comparisons/estimated_" + c.upper() + "_step"+str(s)+".html", title = "estimated_"+c.upper()+"step"+str(s))
            f = figure(title="Estimated colourgrade " +c.upper()+ ": original VS "+str(steps_perc[i])+"%") if useperc else \
                figure(title="Estimated colourgrade " +c.upper()+ ": original VS step"+str(s))
            f.line(df["row"], df[ c + "_avg"], line_color="green", line_width=1, legend="Original")
            f.line(df["row"], df["step"+str(s)+ "_" + c + "_avg"], line_color="red", line_width=1, legend=str(steps_perc[i])+"%" if useperc else "Step"+str(s))
            f.legend.location = "top_left"
            f.yaxis.formatter = NumeralTickFormatter(format="0.00")
            f.xaxis.axis_label = "vineyard rows"
            f.yaxis.axis_label = "%" + c.upper()
            save(f)



def chart_estimates_comparison(folder, id, steps=[], useperc=True):
    """
    Chart statistics by row for estimated berries

    -estimated berries vs originals
    -delta averages vs delta standard deviations
    -delta averages vs delta standard deviations (areas in one chart)

    starting file is the _count_stats_byrow.csv

    """

    if len(steps)>19:
        raise ValueError("number of steps cannot be more than 19")

    #####################################
    ###estimated berries vs originals


    cols = ["estimated_berries_original", "estimated_berries","delta_avg","delta_std" ] + ["step" + str(x) + "_" + "estimated_berries" for x in  steps] \
           + ["step" + str(x) + "_" + "delta_avg" for x in steps] + ["step" + str(x) + "_" + "delta_std" for x in steps]

    path = folder + id + "_count_stats_byrow.csv"
    df = pd.read_csv(path, usecols=cols)


    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries"):
        os.mkdir(folder + "/charts/estimatedberries")

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries/estimatedberries_comparisons"):
        os.mkdir(folder + "/charts/estimatedberries/estimatedberries_comparisons")

    #output chart for original vs no interpolation

    model = linear_model.LinearRegression()
    X, y = df["estimated_berries_original"].values.reshape( df.shape[0],1), df["estimated_berries"].values.reshape( df.shape[0],1)
    model.fit(X, y)
    rsquared = model.score(X, y)

    output_file(folder +"/charts/estimatedberries/estimatedberries_comparisons/estimated_nofilter.html", title = "estimated_nofilter")
    f = figure(title="Estimated berries: original VS no filter")
    f.circle(df["estimated_berries_original"], df["estimated_berries"], fill_color="red", line_color="red")

    f.line(df["estimated_berries_original"], model.predict(df["estimated_berries_original"].values.reshape( df.shape[0],1))[:,0], line_color="blue")

    f.yaxis.formatter = NumeralTickFormatter(format="00")
    f.xaxis.formatter = NumeralTickFormatter(format="00")
    f.xaxis.axis_label = "estimated berries original"
    f.yaxis.axis_label = "estimated berries no_filter"

    citation = Label(x=df["estimated_berries"].min(), y= df["estimated_berries"].max() , x_units='data', y_units='data',
                     text="R-Squared = %.3f"%rsquared, border_line_color='black', border_line_alpha=1.0,
                     background_fill_color="white", background_fill_alpha=1.0)

    f.add_layout(citation)
    save(f)

    steps_perc = [int(1 / int(s) * 100) for s in steps] if useperc else steps

    #iterate steps and output original VS interpolated

    for i,s in enumerate(steps):

        X, y = df["estimated_berries_original"].values.reshape(df.shape[0], 1), \
               df["step" + str(s) + "_" + "estimated_berries"].values.reshape(df.shape[0], 1)
        model.fit(X, y)
        rsquared = model.score(X, y)

        output_file(folder + "/charts/estimatedberries/estimatedberries_comparisons/estimated_step"+str(s)+".html", title="estimated_step"+str(s)+".html")

        f = figure(title="Estimated berries: original VS "+str(steps_perc[i])+"% dataset") if useperc else  \
            figure(title="Estimated berries: original VS step"+str(s))

        f.circle(df["estimated_berries_original"], df["step" + str(s) + "_" +"estimated_berries"], fill_color="red", line_color="red")

        f.line(df["estimated_berries_original"],
               model.predict(df["estimated_berries_original"].values.reshape(df.shape[0], 1))[:, 0], line_color="blue")

        f.yaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.axis_label = "estimated berries original"
        f.yaxis.axis_label = "estimated berries " + str(steps_perc[i]) + "% dataset" if useperc else "estimated berries " +"step " + str(s)

        citation = Label(x=df["estimated_berries"].min(), y=df["estimated_berries"].max(), x_units='data',
                         y_units='data', text="R-Squared = %.3f" % rsquared, border_line_color='black', border_line_alpha=1.0,
                         background_fill_color="white", background_fill_alpha=1.0)

        f.add_layout(citation)
        save(f)

    ###################################################################################
    ##edlta averages vs delta standard deviations

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries/estimatedberries_deltas"):
        os.mkdir(folder + "/charts/estimatedberries/estimatedberries_deltas")

    output_file(folder +"/charts/estimatedberries/estimatedberries_deltas/estimated_deltas_nofilter.html", title = "estimated_deltas_nofilter")
    f = figure(title="Delta average VS Delta std  |  no filter")

    f.circle(df["delta_avg"], df["delta_std"], fill_color="red", line_color="red")

    f.yaxis.formatter = NumeralTickFormatter(format="00")
    f.xaxis.formatter = NumeralTickFormatter(format="00")
    f.xaxis.axis_label = "delta average"
    f.yaxis.axis_label = "delta std"

    save(f)

    #iterate steps and output original VS interpolated

    for i,s in enumerate(steps):

        output_file(folder + "/charts/estimatedberries/estimatedberries_deltas/estimated_step"+str(s)+".html", title="estimated_step"+str(s)+".html")

        f = figure(title="Delta average VS Delta std  |  "+str(steps_perc[i])+"% dataset") if useperc else  \
            figure(title="Delta average VS Delta std  |  step"+str(s))

        f.circle(df["step" + str(s) + "_" +"delta_avg"], df["step" + str(s) + "_" +"delta_std"], fill_color="red", line_color="red")

        f.yaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.axis_label = "delta average"
        f.yaxis.axis_label = "delta std"

        save(f)

    ######show delta areas in one chart

    colors = d3['Category20'][len(steps)+1] if len(steps)>3 else d3['Category20'][4]

    output_file(folder + "/charts/estimatedberries/estimatedberries_deltas/estimated_comparisons.html",
                title="estimated_comparisons.html")
    f = figure(title="Delta average VS Delta std")

    points = df[["delta_avg", "delta_std"]].values
    hull = ConvexHull(points)
    f.patch(points[hull.vertices,0], points[hull.vertices,1], line_alpha=1, fill_alpha=0.0, line_width=2,color=colors[0], legend="No filter")

    for i,s in enumerate(steps):
        points = df[["step" + str(s) + "_" +"delta_avg", "step" + str(s) + "_" +"delta_std"]].values
        hull = ConvexHull(points)
        f.patch(points[hull.vertices,0], points[hull.vertices,1], line_alpha=1, fill_alpha=0.0, line_width=2,color=colors[i+1],
                legend=str(steps_perc[i])+"%" if useperc else "Step"+ str(s))
    f.yaxis.formatter = NumeralTickFormatter(format="0")
    f.xaxis.formatter = NumeralTickFormatter(format="0")
    f.xaxis.axis_label = "delta average"
    f.yaxis.axis_label = "delta std"
    save(f)


def chart_visible_berries(folder,id, steps=[], useperc=True):
    """Chart  visible berries per meter (all points vs interpolated)

    data source is _ERRORSTATISTICS.csv

    :param folder:
    :param id:
    :param steps:
    :param useperc:
    :return:
    """

    cols = ["visible_fruit_per_m"] + [ "step"+str(s)+"visible" for s in steps]

    path = folder + id + "_ERRORSTATISTICS.csv"
    df = pd.read_csv(path, usecols=cols)

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries"):
        os.mkdir(folder + "/charts/estimatedberries")

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts/estimatedberries/estimatedberries_all_comparisons"):
        os.mkdir(folder + "/charts/estimatedberries/estimatedberries_all_comparisons")

    steps_perc = [int(1 / int(s) * 100) for s in steps] if useperc else steps

    model = linear_model.LinearRegression()

    for i,s in enumerate(steps):

        X, y = df["visible_fruit_per_m"].values.reshape(df.shape[0], 1), \
               df["step"+str(s)+"visible"].values.reshape(df.shape[0], 1)
        model.fit(X, y)
        rsquared = model.score(X, y)

        output_file(folder + "/charts/estimatedberries/estimatedberries_all_comparisons/estimated_step"+str(s)+".html", title="estimated_step"+str(s)+".html")

        f = figure(title="Visible berries per meter: original VS "+str(steps_perc[i])+"% dataset") if useperc else  \
            figure(title="Visible berries per meter: original VS step"+str(s))

        f.circle(df["visible_fruit_per_m"], df["step"+str(s)+"visible"], fill_color="red", line_color="red")

        f.line(df["visible_fruit_per_m"],
               model.predict(df["visible_fruit_per_m"].values.reshape(df.shape[0], 1))[:, 0], line_color="blue")

        f.yaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.formatter = NumeralTickFormatter(format="00")
        f.xaxis.axis_label = "Visible berries per meter original"
        f.yaxis.axis_label = "Visible berries per meter " + str(steps_perc[i]) + "% dataset" if useperc else "estimated berries " +"step " + str(s)

        citation = Label(x=df["visible_fruit_per_m"].min(), y=df["step"+str(s)+"visible"].max(), x_units='data',
                         y_units='data', text="R-Squared = %.3f" % rsquared, border_line_color='black', border_line_alpha=1.0,
                         background_fill_color="white", background_fill_alpha=1.0)

        f.add_layout(citation)
        save(f)

def chart_colorgrades(folder,id, colors=[], steps=[], useperc=True):
    """Chart  colorgrade values (all points vs interpolated)

    data source is _ERRORSTATISTICS.csv

    :param folder:
    :param id:
    :param steps:
    :param useperc:
    :return:
    """

    cols = colors + ["step"+str(s)+ c.lower() for s in steps for c in colors]

    path = folder + id + "_ERRORSTATISTICS.csv"
    df = pd.read_csv(path, usecols=cols)

    #make new chart directory if it does not exist
    if not os.path.exists(folder +"/charts"):
        os.mkdir(folder + "/charts")

    # make new chart directory if it does not exist
    if not os.path.exists(folder + "/charts/colorgrades"):
        os.mkdir(folder + "/charts/colorgrades")

    for cl in colors:
        #make new chart directory if it does not exist
        if not os.path.exists(folder +"/charts/colorgrades/colors"+cl.upper()+"_all_comparisons"):
            os.mkdir(folder + "/charts/colorgrades/colors"+cl.upper()+"_all_comparisons")

    steps_perc = [int(1 / int(s) * 100) for s in steps] if useperc else steps

    model = linear_model.LinearRegression()

    for cl in colors:

        for i,s in enumerate(steps):

            X, y = df[cl.lower()].values.reshape(df.shape[0], 1), \
                   df["step"+str(s)+cl.lower()].values.reshape(df.shape[0], 1)
            model.fit(X, y)
            rsquared = model.score(X, y)

            output_file(folder + "/charts/colorgrades/colors"+cl.upper()+"_all_comparisons/color"+cl.upper()+"_step"+str(s)+".html", title="color"+cl.upper()+"_step"+str(s)+".html")

            f = figure(title="Colorgrade"+cl.upper()+": original VS "+str(steps_perc[i])+"% dataset") if useperc else  \
                figure(title="Colorgrade"+cl.upper()+": original VS step"+str(s))

            f.circle(df[cl.lower()], df["step"+str(s)+cl.lower()], fill_color="red", line_color="red")

            f.line(df[cl.lower()],
                   model.predict(df[cl.lower()].values.reshape(df.shape[0], 1))[:, 0], line_color="blue")

            f.yaxis.formatter = NumeralTickFormatter(format="0.000")
            f.xaxis.formatter = NumeralTickFormatter(format="0.000")
            f.xaxis.axis_label = "Colorgrade"+cl.upper()+" original"
            f.yaxis.axis_label = "Colorgrade"+cl.upper()+" " + str(steps_perc[i]) + "% dataset" if useperc else "Colorgrade"+cl.upper()+ " step " + str(s)

            citation = Label(x=df[cl.lower()].min(), y=df["step"+str(s)+cl.lower()].max(), x_units='data',
                             y_units='data', text="R-Squared = %.3f" % rsquared, border_line_color='black', border_line_alpha=1.0,
                             background_fill_color="white", background_fill_alpha=1.0)

            f.add_layout(citation)
            save(f)

def export_estimates_percentages(id,folder, steps = []):
    """
    Open the _count_stats_byrow.csv file and calculate the difference of estimated berries at each step as percentage
    the last row has the average difference in percentage
    output "_count_difference%.csv"
    :param id:
    :param folder:
    :param steps:
    :return:
    """

    cols = ["row", "estimated_berries_original"]
    [cols.append("step"+str(s)+"_estimated_berries") for s in steps]
    [cols.append("step" + str(s) + "_delta_estimates") for s in steps]
    #print(cols)

    path = folder + id + "_count_stats_byrow.csv"
    df = pd.read_csv(path, usecols=cols)

    #recalculate the deltas with negative
    for s in steps:
        df["step"+str(s)+"_delta_estimates"] = df["step"+str(s)+"_estimated_berries"] - df["estimated_berries_original"]

    out = np.zeros((df.shape[0]+1, len(steps) + 2))

    out[:-1,0] = df["row"].values
    out[:-1,1] = df["estimated_berries_original"]

    for i,s in enumerate(steps):
        out[:-1, i+2] = df["step"+str(s)+"_delta_estimates"]*100/df["estimated_berries_original"]
        out[-1,i+2] = np.mean(out[:-1, i+2]) #this is the last row with the averages

    header = "ROW,EXPECTED"
    for s in steps:
        header+=","+str(int(100/s)) + "%"
    #print(header)

    np.savetxt(folder + id + "_count_difference%.csv", out, fmt='%.3f', delimiter=',', newline='\n', header=header, footer='',
           comments='')

if __name__ == "__main__":
    pass

    #steps=[2,3,4,5,6,7,8,9,10]


    #id, stats = berrycolor_workflow_1(steps,"average")
    #id, stats = berrycolor_workflow_1(steps,"average")
    #folder = "/vagrant/code/pysdss/data/output/text/"+id+"/"
    #fileIO.save_object( folder + id + "_statistics", stats)
    #id = "a8aa8edad1c9d457b826f3e156edf0fae"
    #folder = "/vagrant/code/pysdss/data/output/text/" +id + "/"
    #stats = fileIO.load_object(folder + id + "_statistics")
    #test_interpret_result(stats,id,folder)
    #id = "a5d75134d88a843f388af925670d89772"
    #folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    #point = folder + id + "_keep.csv"
    #test_compare_raster_totruth(point,"a0857fd6fbb0e4f6fbacb26b1a984779a", folder,step=steps )
    #test_compare_raster_totruth(point, id, folder,step=steps, clean=[[-1],[0]])


    #id = "a9dd672655abb4188b6d8ac2cc830b3e4"
    #folder = "/vagrant/code/pysdss/data/output/text/workflow1/inversedistancetopower/190517/"+ id + "/"
    #point = folder + id + "_keep.csv"
    #test_compare_raster_totruth(point, id, folder, step=steps, clean=[[-1], [0]])
    #calculate_rmse_row(folder, id)

    #id = "a5d75134d88a843f388af925670d89772"
    #folder = "/vagrant/code/pysdss/data/output/text/workflow1/movingaverage/"+ id + "/"
    #estimate_berries_byrow(folder,id)


    #id = "a9dd672655abb4188b6d8ac2cc830b3e4"
    #folder = "/vagrant/code/pysdss/data/output/text/workflow1/inversedistancetopower/190517/" + id + "/"
    #estimate_berries_byrow(folder,id)
    #calculate_rmse(folder, id)

    #id = "a5d75134d88a843f388af925670d89772"
    #folder = "/vagrant/code/pysdss/data/output/text/workflow1/movingaverage/" + id + "/"
    #estimate_berries_byrow(folder,id)
    #calculate_rmse(folder, id)

    ##output plots
    #id = "a5d75134d88a843f388af925670d89772"
    #folder = "/vagrant/code/pysdss/data/output/text/workflow1/movingaverage/" + id + "/"
    #chart_rmse(folder,id)


    '''
    #############################workflow1 updated
    steps=[2,3,4,5,6,7,8,9,10,20,30,50,100]


    # original worflow with ordered filter and overlapping interpolation
    id, stats = berrycolor_workflow_1(steps,"average")
    folder = "/vagrant/code/pysdss/data/output/text/"+id+"/"
    fileIO.save_object( folder + id + "_statistics", stats)
    test_interpret_result(stats, id, folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth(point, id, folder, step=steps, clean=[[-1], [0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id,step=steps)
    calculate_rmse(folder,id,nsteps=steps)
    chart_rmse(folder, id)

    id, stats = berrycolor_workflow_1(steps)
    folder = "/vagrant/code/pysdss/data/output/text/"+id+"/"
    fileIO.save_object( folder + id + "_statistics", stats)
    test_interpret_result(stats, id, folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth(point, id, folder, step=steps, clean=[[-1], [0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)
    #######################################

    # worflow with random filter (with first_last row point) and overlapping interpolation
    id, stats = berrycolor_workflow_1(steps, "average",random_filter=True, first_last=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    test_interpret_result(stats, id, folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth(point, id, folder, step=steps, clean=[[-1], [0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)

    # worflow with random filter (with first_last row point) and overlapping interpolation
    id, stats = berrycolor_workflow_1(steps, "invdist", random_filter=True, first_last=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    test_interpret_result(stats, id, folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth(point, id, folder, step=steps, clean=[[-1], [0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)
    #######################################


    # workflow with ordered filter  and force interpolation by row
    id, stats = berrycolor_workflow_1(steps, "average",force_interpolation_by_row=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    stats = fileIO.load_object(folder + id + "_statistics")
    test_interpret_result(stats,id,folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth_byrow(point, id, folder,step=steps, clean=[[-1],[0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)

    # workflow with ordered filter  and force interpolation by row
    id, stats = berrycolor_workflow_1(steps, "invdist",force_interpolation_by_row=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    stats = fileIO.load_object(folder + id + "_statistics")
    test_interpret_result(stats,id,folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth_byrow(point, id, folder,step=steps, clean=[[-1],[0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)

    ################################

    # workflow with random filter (with first_last row point)  and force interpolation by row
    id, stats = berrycolor_workflow_1(steps, "average",force_interpolation_by_row=True,random_filter=True, first_last=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    stats = fileIO.load_object(folder + id + "_statistics")
    test_interpret_result(stats,id,folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth_byrow(point, id, folder,step=steps, clean=[[-1],[0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)

    # workflow with random filter (with first_last row point)  and force interpolation by row
    id, stats = berrycolor_workflow_1(steps, "invdist",force_interpolation_by_row=True,random_filter=True, first_last=True)
    folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
    fileIO.save_object(folder + id + "_statistics", stats)
    stats = fileIO.load_object(folder + id + "_statistics")
    test_interpret_result(stats,id,folder)
    point = folder + id + "_keep.csv"
    test_compare_raster_totruth_byrow(point, id, folder,step=steps, clean=[[-1],[0]])
    calculate_rmse_row(folder, id, nsteps=steps)
    estimate_berries_byrow(folder,id, step=steps)
    calculate_rmse(folder, id, nsteps=steps)
    chart_rmse(folder, id)
    '''

    ####################################
    '''
    steps = [2, 3, 4, 5, 6, 7, 8, 9, 10, 20, 30, 50, 100]
    grades=['a','b','c','d']
    id= "a40b9fc7181184bf0b9835f830a677513"
    folder = "/vagrant/code/pysdss/data/output/text/analysis0806_steps2to100/" + id + "/"
    chart_rmse(folder,id,useperc=True)
    chart_estimated_berries(folder,id, steps=steps, useperc=True)
    chart_estimates_comparison(folder, id, steps=steps)
    chart_visible_berries(folder, id, steps=steps)
    chart_colorgrades(folder,id,grades,steps, useperc=True)
    '''
    #######################################KRIGING#################################


    #################################################################
    ''' Compare RMSE for each couple and output a text file with directory infos and camparison charts
    maindir = "/vagrant/code/pysdss/data/output/text/analysys2206"
    #dirs = os.listdir("/vagrant/code/pysdss/data/output/text/analysys2206")
    dirs = [d for d in os.listdir(maindir) if os.path.isdir(os.path.join(maindir, d))]
    paths = [ maindir + "/" + dir + "/" for dir in dirs]
    if not os.path.exists(maindir+"/comparisons"): os.mkdir(maindir+"/comparisons")
    outputinfotext(paths, maindir+"/comparisons") #write all the directory infos to a single file
    compare_rmse_couples(paths, maindir+"/comparisons")
    '''

    # Compare colorgrade RMSE for each couple and output a text file with directory infos and camparison charts

    '''
    maindir = "/vagrant/code/pysdss/data/output/text/analysys2206"
    #dirs = os.listdir("/vagrant/code/pysdss/data/output/text/analysys2206")
    dirs = [d for d in os.listdir(maindir) if os.path.isdir(os.path.join(maindir, d)) and d != "comparisons" and d != "comparisons_colors"]
    paths = [ maindir + "/" + dir + "/" for dir in dirs]
    if not os.path.exists(maindir+"/comparisons_colors"): os.mkdir(maindir+"/comparisons_colors")
    outputinfotext(paths, maindir+"/comparisons_colors") #write all the directory infos to a single file

    colors=["a","b","c","d"]
    for c in colors:
        compare_rmse_couples(paths, maindir+"/comparisons_colors",colorgrade=c)
    '''

    ################################################################




    ######### OUTPUT 32 combinations  (8*4) ############################################

    def analysis_32combinations():

        ###run interpolation for all the combinations and output the charts
        # change the hardcoded variables before running

        #random_filter -> True for random filtering
        #force_interpolation_by_row -> True to force interpolation by vineyard row
        #first_last -> True to take the first and last point of a vineyard row



        ################### AVERAGE     ---- INVDIST #####################

        interp = ["average", "invdist"]
        #interp = ["invdist"]
        grades = ['a', 'b', 'c', 'd']
        steps = [2, 3, 4, 5, 6, 7, 8, 9, 10, 20, 30, 50, 100]
        # random_filter,force_interpolation_by_row, first_last
        combinations = [
            # [False, False, False],
            # [False, False, True],
            # [False, True, False],
            # [False, True, True],
            # [True, False, False], ##########error? step100
            # [True, False, True],
            # [True, True, False],#####error? step100
            # [True, True, True]
        ]

        for i in interp:
            for c in combinations:

                # original worflow with ordered filter and overlapping interpolation
                id, stats = berrycolor_workflow_1(steps=steps, interp=i, random_filter=c[0],
                                                  force_interpolation_by_row=c[1], first_last=c[2])
                folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
                fileIO.save_object(folder + id + "_statistics", stats)
                test_interpret_result(stats, id, folder)
                point = folder + id + "_keep.csv"

                if c[1]:
                    test_compare_raster_totruth_byrow(point, id, folder, steps=steps, clean=[[-1], [0]])
                else:
                    test_compare_raster_totruth(point, id, folder, steps=steps, clean=[[-1], [0]])

                print("calculate rmse")
                calculate_rmse_row(folder, id, steps=steps)
                estimate_berries_byrow(folder, id, steps=steps)
                calculate_rmse(folder, id, steps=steps)
                print("make charts")
                chart_rmse(folder, id)
                chart_estimated_berries(folder, id, steps=steps, useperc=True)
                chart_estimated_colors(folder, id, steps=steps, colorgrades=grades, useperc=True)
                chart_estimates_comparison(folder, id, steps=steps, useperc=True)
                chart_visible_berries(folder, id, steps=steps, useperc=True)
                chart_colorgrades(folder, id, grades, steps, useperc=True)

        ############   kriging ------------- ########################

        # interp = ["ordinary", "simple"]
        interp = ["ordinary"]
        # interp = ["simple"]
        grades = ['a', 'b', 'c', 'd']
        # steps = [5, 6, 7, 8, 9, 10, 20, 30, 50, 100]
        steps = [50]
        counts = ["visible"]
        variogram = "global"
        # random_filter,force_interpolation_by_row, first_last
        combinations = [
            # [False, False, False],
            # [False, False, True],
            # [False, True, False],
            # [False, True, True],
            [True, False, False],  # taking too much time?
            # [True, False, True], #taking too much time?
            # [True, True, False],
            # [True, True, True]
        ]

        for i in interp:
            for c in combinations:

                id, stats = berrycolor_workflow_kriging(steps=steps, counts=counts, interp=i, random_filter=c[0],
                                                        force_interpolation_by_row=c[1], first_last=c[2],
                                                        variogram=variogram,
                                                        rowdirection="x", epsg="32611")

                # id ="a76b7da4d98084ac5863db7adc77eb957"

                folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
                fileIO.save_object(folder + id + "_statistics_" + i, stats)

                colmodel = ["avg", "std", "area", "nozero_area", "zero_area", "visible_count",
                            "visible_avg", "visible_std"]

                # stats = fileIO.load_object(folder + id + "_statistics_" + i)
                test_interpret_result(stats, id, folder, colmodel)

                point = folder + id + "_keep.csv"
                if c[1]:
                    test_compare_raster_totruth_byrow(point, id, folder, steps=steps, clean=[[-1], [0]], counts=counts,
                                                      suffix="_" + i + variogram)
                else:
                    test_compare_raster_totruth(point, id, folder, steps=steps, clean=[[-1], [0]], counts=counts,
                                                suffix="_" + i + variogram)

                print("calculate rmse")
                calculate_rmse_row(folder, id, steps=steps)
                estimate_berries_byrow(folder, id, steps=steps)
                calculate_rmse(folder, id, steps=steps)
                print("make charts")
                chart_rmse(folder, id)
                chart_estimated_berries(folder, id, steps=steps, useperc=True)
                chart_estimated_colors(folder, id, steps=steps, colorgrades=grades, useperc=True)
                chart_estimates_comparison(folder, id, steps=steps, useperc=True)
                chart_visible_berries(folder, id, steps=steps, useperc=True)
                chart_colorgrades(folder, id, grades, steps, useperc=True)


    # analysis_32combinations()


    ##############  publish docs files with the rsme comparisons images in tables
    chartpath = "/vagrant/code/pysdss/data/output/text/analysys2206/comparisons/charts"
    if not os.path.exists(chartpath + "/docs"): os.mkdir(chartpath + "/docs")

    ################################################

    ##check moving average
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingFFF", "moving"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["movingFFF", "moving"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingTT", "movingFT"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["movingTT", "movingFT"])


    ##check inverse distance
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseFFF", "inverse"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseTT", "inverseFT"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["inverseFFF", "inverse"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["inverseTT", "inverseFT"])


    ###140717 additional comparisons  end of section 1.7.1
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseFT", "inverseFF"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["inverseFT", "inverseFF"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingFT", "inverseFF"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["movingFT", "inverseFF"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseTF", "inverseFF"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["inverseTF", "inverseFF"])

    '''
    ##check ordinary kriging
    publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["ordinaryFFF", "ordinary"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["ordinaryTT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["ordinaryFFF", "ordinary"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["ordinaryTT", "ordinaryFT"])
    '''
    '''
    ## check simple kriging
    publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["simpleFFF", "simple"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["simpleTT", "simpleFT"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["simpleFFF", "simple"])
    publish_rsme_comparisons(chartpath,chartpath + "/docs","estimates", ["simpleTT", "simpleFT"])
    '''

    '''
    ## from the above i found that ordered filter restricted to same row gives better results
    ## now i can compare ordered/same row for the 4 interpolation methods (10 comparisons * 2)
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "movingFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["simpleFT", "simpleFT"])
    '''

    '''
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["movingFT", "movingFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["movingFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["movingFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["movingFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["inverseFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["inverseFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["inverseFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["ordinaryFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["ordinaryFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "estimates", ["simpleFT", "simpleFT"])
    '''

    #eporting estimated berries tables for interpolations found after the analysis of the above rmse
    rootdir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    #dirs=["a066f6d0e9a0d49d2999c3714c5f9ed34","a1a126f124b7e4f5e99d4ddc553f71725","a37809e20381e49778efa56222f9bcbce",
    #"a3f8537dee1324c088bdc84562c6c4f0c","a923f00d5dbdc4d988b128377c92dc19a","a97304dec81854237a9cd9523e28daf52",
    #"aa2de2f54ca16491fa46e52eb9c2b3717","aff8a205a38f34bb3bc2cd177bed15dc4"]
    #dirs=["ab0d16e6152d94383bdf5c6dba3a45dea","ac5eaa885603f4d83ad939623bae6d116"]
    #outpath =  rootdir +  "/comparisons/charts/docs/"
    #export_estimated_berries_tables(rootdir, dirs, outpath)



    ##############  publish docs files with the rsme comparisons images in tables for color grades
    #chartpath = "/vagrant/code/pysdss/data/output/text/analysys2206/comparisons_colors/charts"
    #if not os.path.exists(chartpath + "/docs"): os.mkdir(chartpath + "/docs")

    ##check moving average
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingFFF", "moving"])
    #######publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingTT", "movingFT"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["movingFF", "movingFT"])

    ##check inverse distance
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseFFF", "inverse"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["inverseTT", "inverseFT"])


    ##check ordinary kriging
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["ordinaryFFF", "ordinary"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["ordinaryTT", "ordinaryFT"])
    #publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFF", "ordinaryFT"])

    ## check simple kriging
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["simpleFFF", "simple"])
    #publish_rsme_comparisons(chartpath,chartpath + "/docs","avg", ["simpleTT", "simpleFT"])
    #publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["simpleFF", "simpleFT"])



    ## from the above i found that ordered filter restricted to same row gives better results but also ordered with free
    # filtrer can give better results
    ## now i can compare ordered/same row for the 4 interpolation methods (10 comparisons * 2) and then ordered/same row
    '''
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "movingFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "inverseFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFT", "simpleFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFT", "ordinaryFT"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["simpleFT", "simpleFT"])
    '''
    '''
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFF", "movingFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFF", "inverseFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFF", "ordinaryFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["movingFF", "simpleFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFF", "inverseFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFF", "ordinaryFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["inverseFF", "simpleFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFF", "simpleFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["ordinaryFF", "ordinaryFF"])
    publish_rsme_comparisons(chartpath, chartpath + "/docs", "avg", ["simpleFF", "simpleFF"])      
    '''

    ## i still need to export the comparison charts! i do only for the ordered/only rows

    ''''
    rootdir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    dirs=["a066f6d0e9a0d49d2999c3714c5f9ed34","a923f00d5dbdc4d988b128377c92dc19a","a97304dec81854237a9cd9523e28daf52",
    "aff8a205a38f34bb3bc2cd177bed15dc4"]
    paths = [rootdir+d+"/" for d in dirs]
    for i,folder in enumerate(paths):
        chart_estimated_colors(folder, dirs[i], steps=[2,3,4,5,6,7,8,9,10,20,30,50,100], colorgrades=["a", "b", "c", "d"], useperc=True)

    dirs = ["a1a126f124b7e4f5e99d4ddc553f71725", "a37809e20381e49778efa56222f9bcbce",
    "a3f8537dee1324c088bdc84562c6c4f0c","aa2de2f54ca16491fa46e52eb9c2b3717"]
    paths = [rootdir+d+"/" for d in dirs]
    for i,folder in enumerate(paths):
        chart_estimated_colors(folder, dirs[i], steps=[5,6,7,8,9,10,20,30,50,100], colorgrades=["a", "b", "c", "d"], useperc=True)
    '''


    #esporting esrimated colorgrades tables for interpolations found after the analysis of the above rmse
    #rootdir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    #dirs=["a066f6d0e9a0d49d2999c3714c5f9ed34","a1a126f124b7e4f5e99d4ddc553f71725","a37809e20381e49778efa56222f9bcbce",
    #"a3f8537dee1324c088bdc84562c6c4f0c","a923f00d5dbdc4d988b128377c92dc19a","a97304dec81854237a9cd9523e28daf52",
    #"aa2de2f54ca16491fa46e52eb9c2b3717","aff8a205a38f34bb3bc2cd177bed15dc4"]
    #outpath =  rootdir +  "/comparisons_colors/charts/docs/"
    #export_estimated_colors_tables(rootdir, dirs, outpath)



    ###################################################################
    ############################  TESTS ###############################
    ###################################################################

    def kriger_tests():
        steps = [5, 6, 7, 8, 9, 10, 20, 30, 50, 100]
        steps = [50, 100]
        grades = ['a', 'b', 'c', 'd']

        interp = "simple"
        variogram = "global"

        id, stats = berrycolor_workflow_kriging(steps=steps, interp=interp, random_filter=False, first_last=False, force_interpolation_by_row=False,
                             variogram = variogram,rowdirection="x", epsg="32611")

        folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
        fileIO.save_object(folder + id + "_statistics_"+ interp, stats)

        ##############################

        #id = "a7aa8d3e709ce42a4ba17ebee02daf66f"
        folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"
        #folder = "/vagrant/code/pysdss/data/output/text/analysys1906_krigingsteps5to100/" + id + "/"

        colmodel = ["avg", "std", "area", "nozero_area", "zero_area", "visible_count",
                    "visible_avg", "visible_std"]

        stats = fileIO.load_object(folder + id + "_statistics_" + interp)
        test_interpret_result(stats, id, folder, colmodel)

        print("comparing to points")
        point = folder + id + "_keep.csv"

        test_compare_raster_totruth(point, id, folder, steps=steps, clean=[[-1], [0]], counts=["visible"],
                                         suffix="_" + interp + variogram)

        print("calculate rmse")
        calculate_rmse_row(folder, id, steps=steps)
        estimate_berries_byrow(folder, id, steps=steps)
        calculate_rmse(folder, id, steps=steps)
        print("make charts")
        chart_rmse(folder, id)
        chart_estimated_berries(folder, id, steps=steps, useperc=True)
        chart_estimates_comparison(folder, id, steps=steps)
        chart_visible_berries(folder, id, steps=steps)
        chart_colorgrades(folder, id, grades, steps, useperc=True)

    #kriger_tests()

    def comparematrices_tests():
        from osgeo import gdal_array as gdar

        d = gdal.Open(
            "/vagrant/code/pysdss/data/output/text/afc3e0829d87e4aa19212923b3f44e519/afc3e0829d87e4aa19212923b3f44e519_d.tiff")
        band = d.GetRasterBand(1)

        g = d.GetGeoTransform()
        print(d)

        e = gdal.Open(
            "/vagrant/code/pysdss/data/output/text/afc3e0829d87e4aa19212923b3f44e519/afc3e0829d87e4aa19212923b3f44e519_d_step50_ordinaryglobal.tiff")
        band1 = e.GetRasterBand(1)
        e.GetGeoTransform()

        px = gdar.BandReadAsArray(band, 0, 0, band.XSize, band.YSize)
        px1 = gdar.BandReadAsArray(band1, 0, 0, band1.XSize, band1.YSize)

        pass


    # comparematrices_tests()


    def kriger_byrow_tests():
        # steps = [5,6,7,8,9,10,20,30,50,100]
        steps = [50, 100]
        grades = ['a', 'b', 'c', 'd']
        counts = ["visible"]
        interp = "ordinary"
        variogram = "global"

        id, stats = berrycolor_workflow_kriging(steps=steps, colorgrades = grades, counts = counts,
                             interp=interp, random_filter=False, first_last=False, force_interpolation_by_row=True,
                             variogram = variogram,rowdirection="x", epsg="32611")


        #id = "a54c4dd3232e64ab38dfcc8ad47103065"
        folder = "/vagrant/code/pysdss/data/output/text/" + id + "/"

        fileIO.save_object(folder + id + "_statistics_"+ interp, stats)



        colmodel = ["avg", "std", "area", "nozero_area", "zero_area", "visible_count",
                    "visible_avg", "visible_std"]

        stats = fileIO.load_object(folder + id + "_statistics_" + interp)

        test_interpret_result(stats, id, folder, colmodel)

        point = folder + id + "_keep.csv"
        test_compare_raster_totruth_byrow(point, id, folder, steps=steps, clean=[[-1], [0]], counts=counts,
                                               suffix="_" + interp + variogram)

        print("calculate rmse")
        calculate_rmse_row(folder, id, steps=steps)
        estimate_berries_byrow(folder, id, steps=steps)
        calculate_rmse(folder, id, steps=steps)
        print("make charts")
        chart_rmse(folder, id)
        chart_estimated_berries(folder, id, steps=steps, useperc=True)
        chart_estimates_comparison(folder, id, steps=steps)
        chart_visible_berries(folder, id, steps=steps)
        chart_colorgrades(folder, id, grades, steps, useperc=True)


    #kriger_byrow_tests()


    def outinfos_tests():
        """
        test output infos
        :return:
        """
        steps = [5, 6, 7, 8, 9, 10, 20, 30, 50, 100]
        grades = ['a', 'b', 'c', 'd']
        counts = ["visible"]

        interp = "simple"
        variogram = "global"

        berrycolor_workflow_kriging(steps=steps, counts=counts,interp=interp, random_filter=False, first_last=False,
                                                force_interpolation_by_row=False,
                                                variogram=variogram, rowdirection="x", epsg="32611")

        interp = "ordinary"
        variogram = "global"
        berrycolor_workflow_kriging(steps=steps, counts=counts,interp=interp, random_filter=True, first_last=True,
                                                force_interpolation_by_row=True,
                                                variogram=variogram, rowdirection="x", epsg="32611")

        berrycolor_workflow_1(steps=[2, 3, 4], interp="invdist", random_filter=False, first_last=False,
                              force_interpolation_by_row=False, rowdirection="x")

        berrycolor_workflow_1(steps=[2, 3, 4], interp="average", random_filter=False, first_last=False,
                              force_interpolation_by_row=False, rowdirection="x")
    #outinfos_tests()



    def export_word_test(outdir, img):
        """
        just testing creating and exporting word file with text and table with pictures
        :return:
        """
        import docx
        from docx import Document
        from docx.shared import Cm

        document = Document()


        document.add_heading('The REAL meaning of the universe')
        paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
        prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

        document.add_page_break()

        document.add_heading('The role of dolphins', level=2)

        document.add_page_break()

        table = document.add_table(rows=2, cols=2)

        table.style = 'LightShading-Accent1'

        cell = table.cell(0, 0)
        cell.text = 'parrot, possibly dead'

        row = table.rows[1]
        row.cells[0].text = 'Foo bar to you.'
        row.cells[1].text = 'And a hearty foo bar to you too sir!'


        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img, width=Cm(7.5))

        row_count = len(table.rows)
        col_count = len(table.columns)

        print(row_count,col_count)

        row = table.add_row()
        row_count = len(table.rows)
        print(row_count,col_count)

        document.add_picture(img, width=Cm(7.5))
        document.save(outdir+'/test.docx')

    #export_word_test("/vagrant/code/pysdss/data/output/text","/vagrant/code/pysdss/data/output/text/analysys2206/comparisons/charts/delta_avg_rmse_comparison0_simpleTFT_movingTFF.png"  )


    def estimate_colors_byrow_test():

        id = "a80749789a6134142a654caa0090d7a8f"

        maindir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
        folder =  maindir + id + "/"
        dirs = [d for d in os.listdir(maindir) if os.path.isdir(os.path.join(maindir, d)) and d != "comparisons" and d != "comparisons_colors"]
        paths = [maindir + "/" + dir + "/" for dir in dirs]
        #print(paths)
        for d in dirs:
            folder = maindir + d + "/"
            print(folder)
            info = utils.json_io(folder + "/info.json", direction="in")
            estimate_colors_byrow(folder, d, steps = info["steps"], colorgrades = info["colorgrades"])
            calculate_rmse_color(folder, d, steps=info["steps"], colorgrades=info["colorgrades"])
            chart_rmse_color(folder, d, colorgrades=info["colorgrades"], useperc=True)
    #estimate_colors_byrow_test()



    ########################################


    #ordered, byrows
    '''ids = ["a923f00d5dbdc4d988b128377c92dc19a","aff8a205a38f34bb3bc2cd177bed15dc4",
           "a97304dec81854237a9cd9523e28daf52","a066f6d0e9a0d49d2999c3714c5f9ed34"]
    maindir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    for id in ids:
        print(id)
        folder = maindir + id + "/"
        export_estimates_percentages(id,folder, steps = [2,3,4,5,6,7,8,9,10,20,30,50,100])'''

    #ordered, free
    '''ids = ["a2f718fbced16402a8774b750da18052c","a80749789a6134142a654caa0090d7a8f",
           "ab0d16e6152d94383bdf5c6dba3a45dea","ac5eaa885603f4d83ad939623bae6d116"]
    maindir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    for id in ids:
        print(id)
        folder = maindir + id + "/"
        export_estimates_percentages(id,folder, steps = [2,3,4,5,6,7,8,9,10,20,30,50,100])'''
    '''
    ids = ["ab0d16e6152d94383bdf5c6dba3a45dea", "ac5eaa885603f4d83ad939623bae6d116"]
    maindir = "/vagrant/code/pysdss/data/output/text/analysys2206/"
    for id in ids:
        print(id)
        folder = maindir + id + "/"
        export_estimates_percentages(id,folder, steps = [2,3,4,5,6,7,8,9,10,20,30,50,100])
    '''